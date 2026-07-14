"""
Backend regression tests for the NEW features added in this iteration:
- Documents library (upload, list, library-by-role, download, send, delete)
- Blogs / News (draft/publish, public list & detail, upload-image, delete)
- Calendar (ICS public feed, Word .docx import preview + bulk create)

Also lightly re-verifies existing features (auth, broadcast, notifications, enquiries)
so the existing regression surface still holds.

Runs against REACT_APP_BACKEND_URL (Kubernetes ingress).
"""
import io
import os
import time
import pytest
import requests
from datetime import datetime, timedelta

BASE_URL = os.environ.get("REACT_APP_BACKEND_URL", "").rstrip("/")
assert BASE_URL, "REACT_APP_BACKEND_URL must be set"

ADMIN_EMAIL = "admin@1471horwich.org.uk"
ADMIN_PASSWORD = "Squadron1471!"
CADET_EMAIL = "cadet@1471horwich.org.uk"
CADET_PASSWORD = "Cadet1471!"
PARENT_EMAIL = "parent@1471horwich.org.uk"
PARENT_PASSWORD = "Parent1471!"


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------
def _login(email: str, password: str) -> str:
    r = requests.post(f"{BASE_URL}/api/auth/login",
                      json={"email": email, "password": password}, timeout=30)
    assert r.status_code == 200, f"login failed for {email}: {r.status_code} {r.text}"
    return r.json()["access_token"]


@pytest.fixture(scope="module")
def admin_token():
    return _login(ADMIN_EMAIL, ADMIN_PASSWORD)


@pytest.fixture(scope="module")
def cadet_token():
    return _login(CADET_EMAIL, CADET_PASSWORD)


@pytest.fixture(scope="module")
def parent_token():
    return _login(PARENT_EMAIL, PARENT_PASSWORD)


def _auth(token: str):
    return {"Authorization": f"Bearer {token}"}


# ---------------------------------------------------------------------------
# Documents
# ---------------------------------------------------------------------------
class TestDocuments:
    """Document library: upload/list/library/download/send/delete."""

    created_ids = []

    def test_login_admin(self, admin_token):
        assert isinstance(admin_token, str) and len(admin_token) > 10

    def test_upload_document_visible_to_cadet(self, admin_token):
        files = {"file": ("TEST_cadet_doc.txt", b"cadet pack contents", "text/plain")}
        data = {"title": "TEST Cadet Pack", "category": "TEST_Cat_A",
                "visible_roles": "cadet"}
        r = requests.post(f"{BASE_URL}/api/documents", data=data, files=files,
                          headers=_auth(admin_token), timeout=30)
        assert r.status_code == 200, r.text
        d = r.json()
        assert d["title"] == "TEST Cadet Pack"
        assert d["category"] == "TEST_Cat_A"
        assert d["filename"] == "TEST_cadet_doc.txt"
        assert d["visible_roles"] == ["cadet"]
        assert "gridfs_id" not in d
        assert "_id" not in d
        TestDocuments.created_ids.append(d["id"])
        # Verify persisted via list
        r2 = requests.get(f"{BASE_URL}/api/documents", headers=_auth(admin_token), timeout=30)
        assert r2.status_code == 200
        assert any(x["id"] == d["id"] for x in r2.json())

    def test_upload_document_visible_to_parent_only(self, admin_token):
        files = {"file": ("TEST_parent.pdf", b"%PDF-1.4 fake parent doc", "application/pdf")}
        data = {"title": "TEST Parent Only", "category": "TEST_Cat_B",
                "visible_roles": "parent"}
        r = requests.post(f"{BASE_URL}/api/documents", data=data, files=files,
                          headers=_auth(admin_token), timeout=30)
        assert r.status_code == 200
        TestDocuments.created_ids.append(r.json()["id"])

    def test_upload_document_staff_only(self, admin_token):
        files = {"file": ("TEST_staff.txt", b"staff eyes only", "text/plain")}
        # No roles -> not visible to any member role
        data = {"title": "TEST Staff Only", "category": "TEST_Staff",
                "visible_roles": ""}
        r = requests.post(f"{BASE_URL}/api/documents", data=data, files=files,
                          headers=_auth(admin_token), timeout=30)
        assert r.status_code == 200
        d = r.json()
        assert d["visible_roles"] == []
        TestDocuments.created_ids.append(d["id"])

    def test_library_for_cadet_includes_cadet_doc_only(self, cadet_token):
        r = requests.get(f"{BASE_URL}/api/documents/library",
                         headers=_auth(cadet_token), timeout=30)
        assert r.status_code == 200
        docs = r.json()
        titles = [d["title"] for d in docs]
        assert "TEST Cadet Pack" in titles
        assert "TEST Parent Only" not in titles
        assert "TEST Staff Only" not in titles

    def test_library_for_parent_includes_parent_doc(self, parent_token):
        r = requests.get(f"{BASE_URL}/api/documents/library",
                         headers=_auth(parent_token), timeout=30)
        assert r.status_code == 200
        titles = [d["title"] for d in r.json()]
        assert "TEST Parent Only" in titles
        assert "TEST Cadet Pack" not in titles

    def test_download_public(self):
        # Download endpoint is public (no auth) — used by email link recipients.
        doc_id = TestDocuments.created_ids[0]
        r = requests.get(f"{BASE_URL}/api/documents/{doc_id}/download", timeout=30)
        assert r.status_code == 200
        assert r.content == b"cadet pack contents"
        assert "text/plain" in r.headers.get("content-type", "")

    def test_patch_document_updates_roles(self, admin_token):
        doc_id = TestDocuments.created_ids[0]
        r = requests.patch(f"{BASE_URL}/api/documents/{doc_id}",
                           json={"title": "TEST Cadet Pack v2",
                                 "category": "TEST_Cat_A",
                                 "visible_roles": ["cadet", "parent"]},
                           headers=_auth(admin_token), timeout=30)
        assert r.status_code == 200
        assert r.json()["visible_roles"] == ["cadet", "parent"]

    def test_send_document_to_parents_creates_notifications(self, admin_token, parent_token):
        doc_id = TestDocuments.created_ids[1]  # TEST Parent Only
        # Snapshot parent notifications count before
        r0 = requests.get(f"{BASE_URL}/api/notifications", headers=_auth(parent_token), timeout=30)
        before = len(r0.json())
        payload = {
            "audience": {"mode": "roles", "roles": ["parent"], "user_ids": []},
            "channels": ["dashboard"],   # skip email to keep the test fast
            "message": "Please read.",
            "base_url": BASE_URL,
        }
        r = requests.post(f"{BASE_URL}/api/documents/{doc_id}/send",
                          json=payload, headers=_auth(admin_token), timeout=60)
        assert r.status_code == 200, r.text
        result = r.json()
        assert result["recipients"] >= 1
        assert result["dashboard_delivered"] >= 1
        # Parent should now see a notification with a download link
        r1 = requests.get(f"{BASE_URL}/api/notifications", headers=_auth(parent_token), timeout=30)
        assert r1.status_code == 200
        after = r1.json()
        assert len(after) > before
        newest = after[0]
        assert newest["kind"] == "document"
        assert newest.get("link", "").endswith(f"/api/documents/{doc_id}/download")
        assert newest.get("link_label") == "Download document"

    def test_send_document_missing_recipients_returns_400(self, admin_token):
        doc_id = TestDocuments.created_ids[0]
        payload = {"audience": {"mode": "users", "user_ids": []},
                   "channels": ["dashboard"], "message": "", "base_url": BASE_URL}
        r = requests.post(f"{BASE_URL}/api/documents/{doc_id}/send",
                          json=payload, headers=_auth(admin_token), timeout=30)
        assert r.status_code == 400

    def test_cadet_cannot_upload(self, cadet_token):
        files = {"file": ("bad.txt", b"x", "text/plain")}
        data = {"title": "nope", "category": "x", "visible_roles": "cadet"}
        r = requests.post(f"{BASE_URL}/api/documents", data=data, files=files,
                          headers=_auth(cadet_token), timeout=30)
        assert r.status_code in (401, 403)

    def test_delete_documents_cleanup(self, admin_token):
        for did in TestDocuments.created_ids:
            r = requests.delete(f"{BASE_URL}/api/documents/{did}",
                                headers=_auth(admin_token), timeout=30)
            assert r.status_code == 200
        # Verify at least one is gone
        r = requests.get(f"{BASE_URL}/api/documents/{TestDocuments.created_ids[0]}/download",
                         timeout=30)
        assert r.status_code == 404


# ---------------------------------------------------------------------------
# Blogs / News
# ---------------------------------------------------------------------------
class TestBlogs:
    created_ids = []
    published_slug = None
    draft_slug = None

    def test_create_published_post(self, admin_token):
        payload = {"title": "TEST Squadron News Post",
                   "excerpt": "TEST excerpt line",
                   "body": "This is the *body* of the news post.",
                   "cover_image_url": "",
                   "images": [],
                   "status": "published"}
        r = requests.post(f"{BASE_URL}/api/blogs", json=payload,
                          headers=_auth(admin_token), timeout=30)
        assert r.status_code == 200, r.text
        b = r.json()
        assert b["status"] == "published"
        assert b["published_at"]
        assert b["slug"]
        assert "_id" not in b
        TestBlogs.created_ids.append(b["id"])
        TestBlogs.published_slug = b["slug"]

    def test_create_draft_post(self, admin_token):
        payload = {"title": "TEST Draft Post",
                   "excerpt": "hidden",
                   "body": "draft only",
                   "status": "draft"}
        r = requests.post(f"{BASE_URL}/api/blogs", json=payload,
                          headers=_auth(admin_token), timeout=30)
        assert r.status_code == 200
        b = r.json()
        assert b["status"] == "draft"
        assert b.get("published_at") in (None, "")
        TestBlogs.created_ids.append(b["id"])
        TestBlogs.draft_slug = b["slug"]

    def test_public_list_hides_draft(self):
        r = requests.get(f"{BASE_URL}/api/public/blogs", timeout=30)
        assert r.status_code == 200
        slugs = [p["slug"] for p in r.json()]
        assert TestBlogs.published_slug in slugs
        assert TestBlogs.draft_slug not in slugs

    def test_public_detail_returns_published(self):
        r = requests.get(f"{BASE_URL}/api/public/blogs/{TestBlogs.published_slug}",
                         timeout=30)
        assert r.status_code == 200
        b = r.json()
        assert b["title"] == "TEST Squadron News Post"
        assert b["body"].startswith("This is the")

    def test_public_detail_draft_returns_404(self):
        r = requests.get(f"{BASE_URL}/api/public/blogs/{TestBlogs.draft_slug}",
                         timeout=30)
        assert r.status_code == 404

    def test_patch_publish_draft(self, admin_token):
        # Note: PATCH uses the same BlogCreate schema as POST, so title+body
        # must be supplied. This is how the frontend BlogAdminPanel calls it.
        draft_id = TestBlogs.created_ids[1]
        r = requests.patch(f"{BASE_URL}/api/blogs/{draft_id}",
                           json={"title": "TEST Draft Post",
                                 "excerpt": "hidden",
                                 "body": "draft only, now published",
                                 "status": "published"},
                           headers=_auth(admin_token), timeout=30)
        assert r.status_code == 200
        b = r.json()
        assert b["status"] == "published"
        assert b.get("published_at")

    def test_cadet_cannot_create_blog(self, cadet_token):
        r = requests.post(f"{BASE_URL}/api/blogs",
                          json={"title": "no", "body": "no"},
                          headers=_auth(cadet_token), timeout=30)
        assert r.status_code in (401, 403)

    def test_delete_blogs_cleanup(self, admin_token):
        for bid in TestBlogs.created_ids:
            r = requests.delete(f"{BASE_URL}/api/blogs/{bid}",
                                headers=_auth(admin_token), timeout=30)
            assert r.status_code == 200


# ---------------------------------------------------------------------------
# Calendar: ICS feed + Word import
# ---------------------------------------------------------------------------
class TestCalendar:
    created_event_ids = []

    def test_ics_feed_is_valid(self):
        r = requests.get(f"{BASE_URL}/api/calendar/events.ics", timeout=30)
        assert r.status_code == 200
        ctype = r.headers.get("content-type", "")
        assert "text/calendar" in ctype
        text = r.text
        assert text.startswith("BEGIN:VCALENDAR")
        assert "END:VCALENDAR" in text
        # Should list at least one VEVENT from seed data
        assert "BEGIN:VEVENT" in text
        assert "END:VEVENT" in text

    def test_import_docx_preview(self, admin_token):
        # Build a tiny .docx with a date+activity table via python-docx
        try:
            from docx import Document
        except Exception as e:
            pytest.skip(f"python-docx not available: {e}")
        doc = Document()
        doc.add_heading("Training Programme", level=1)
        table = doc.add_table(rows=4, cols=2)
        table.rows[0].cells[0].text = "Date"
        table.rows[0].cells[1].text = "Activity"
        table.rows[1].cells[0].text = "Thursday 5 June 2026"
        table.rows[1].cells[1].text = "Drill night"
        table.rows[2].cells[0].text = "Monday 9 June 2026"
        table.rows[2].cells[1].text = "First aid"
        table.rows[3].cells[0].text = "Thursday 12 June 2026"
        table.rows[3].cells[1].text = "Fieldcraft"
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        files = {"file": ("TEST_programme.docx", buf.getvalue(),
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        r = requests.post(f"{BASE_URL}/api/events/import-docx", files=files,
                          headers=_auth(admin_token), timeout=60)
        assert r.status_code == 200, r.text
        preview = r.json()
        assert preview["count"] == 3
        titles = [e["title"] for e in preview["events"]]
        assert "Drill night" in titles
        assert "First aid" in titles
        assert "Fieldcraft" in titles
        # And each has a proper ISO start
        for ev in preview["events"]:
            assert ev["start"].startswith("2026-06-")
            assert "date_label" in ev

    def test_import_events_bulk_create(self, admin_token):
        payload = {"events": [
            {"title": "TEST_ImportEvt A",
             "start": "2026-07-15T19:00:00",
             "end": "2026-07-15T21:30:00",
             "description": "",
             "location": "Squadron HQ"},
            {"title": "TEST_ImportEvt B",
             "start": "2026-07-22T19:00:00",
             "end": "2026-07-22T21:30:00",
             "description": "",
             "location": "Squadron HQ"},
        ]}
        r = requests.post(f"{BASE_URL}/api/events/import", json=payload,
                          headers=_auth(admin_token), timeout=30)
        assert r.status_code == 200
        assert r.json()["created"] == 2

        # Find & remember to delete after
        r2 = requests.get(f"{BASE_URL}/api/events", headers=_auth(admin_token), timeout=30)
        assert r2.status_code == 200
        for e in r2.json():
            if e["title"].startswith("TEST_ImportEvt"):
                TestCalendar.created_event_ids.append(e["id"])
        assert len(TestCalendar.created_event_ids) >= 2

    def test_ics_feed_includes_newly_imported_events(self):
        r = requests.get(f"{BASE_URL}/api/calendar/events.ics", timeout=30)
        assert r.status_code == 200
        assert "TEST_ImportEvt A" in r.text or "TEST_ImportEvt B" in r.text

    def test_import_docx_rejects_non_docx(self, admin_token):
        files = {"file": ("bad.txt", b"not a docx", "text/plain")}
        r = requests.post(f"{BASE_URL}/api/events/import-docx", files=files,
                          headers=_auth(admin_token), timeout=30)
        assert r.status_code == 400

    def test_cleanup_imported_events(self, admin_token):
        for eid in TestCalendar.created_event_ids:
            r = requests.delete(f"{BASE_URL}/api/events/{eid}",
                                headers=_auth(admin_token), timeout=30)
            assert r.status_code in (200, 204)


# ---------------------------------------------------------------------------
# Regression on existing endpoints
# ---------------------------------------------------------------------------
class TestRegression:
    def test_auth_me(self, admin_token):
        r = requests.get(f"{BASE_URL}/api/auth/me", headers=_auth(admin_token), timeout=30)
        assert r.status_code == 200
        assert r.json()["email"] == ADMIN_EMAIL

    def test_notifications_endpoint(self, cadet_token):
        r = requests.get(f"{BASE_URL}/api/notifications",
                         headers=_auth(cadet_token), timeout=30)
        assert r.status_code == 200
        assert isinstance(r.json(), list)

    def test_broadcast_dashboard_only(self, admin_token, cadet_token):
        payload = {"title": "TEST Broadcast",
                   "body": "hello cadets",
                   "audience": {"mode": "roles", "roles": ["cadet"]},
                   "channels": ["dashboard"]}
        r = requests.post(f"{BASE_URL}/api/broadcast", json=payload,
                          headers=_auth(admin_token), timeout=30)
        assert r.status_code == 200
        assert r.json()["dashboard_delivered"] >= 1
        # cadet sees it
        r2 = requests.get(f"{BASE_URL}/api/notifications",
                          headers=_auth(cadet_token), timeout=30)
        assert any(n["title"] == "TEST Broadcast" for n in r2.json())

    def test_enquiries_tracker(self, admin_token):
        r = requests.get(f"{BASE_URL}/api/enquiries/tracker",
                         headers=_auth(admin_token), timeout=30)
        assert r.status_code == 200
        body = r.json()
        buckets = body.get("buckets", {})
        assert "now" in buckets and "september" in buckets and "future" in buckets

    def test_events_list(self, admin_token):
        r = requests.get(f"{BASE_URL}/api/events", headers=_auth(admin_token), timeout=30)
        assert r.status_code == 200
        assert isinstance(r.json(), list)
