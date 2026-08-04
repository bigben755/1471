from dotenv import load_dotenv
from pathlib import Path

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

import os
import re
import io
import json
import html
import base64
import logging
import uuid
import asyncio
from calendar import monthrange
from datetime import datetime, timezone, timedelta, date
from typing import Dict, List, Optional

import jwt
import bcrypt
import httpx
from bson import ObjectId
from fastapi import FastAPI, APIRouter, HTTPException, Depends, UploadFile, File, Form, BackgroundTasks
from fastapi.responses import Response
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient, AsyncIOMotorGridFSBucket
from pydantic import BaseModel, Field, EmailStr, ConfigDict

try:
    import resend
except Exception:  # pragma: no cover
    resend = None

# ---------------------------------------------------------------------------
# Setup
# ---------------------------------------------------------------------------
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]
fs = AsyncIOMotorGridFSBucket(db)

JWT_SECRET = os.environ['JWT_SECRET']
JWT_ALGORITHM = "HS256"
ADMIN_EMAIL = os.environ.get('ADMIN_EMAIL', 'admin@example.com')
ADMIN_PASSWORD = os.environ.get('ADMIN_PASSWORD', 'admin123')
CADET_DEFAULT_PASSWORD = os.environ.get('CADET_DEFAULT_PASSWORD', 'Squadron123!')
DEFAULT_EVENT_LINK = os.environ.get('DEFAULT_EVENT_LINK', 'https://cadets.bader.mod.uk/')
RESEND_API_KEY = os.environ.get('RESEND_API_KEY', '').strip()
SENDER_EMAIL = os.environ.get('SENDER_EMAIL', 'onboarding@resend.dev')
NOTIFY_EMAIL = os.environ.get('NOTIFY_EMAIL', ADMIN_EMAIL)
SMS_RESET_EMAIL = os.environ.get('SMS_RESET_EMAIL', 'oc.1471@rafac.mod.gov.uk').strip()

# Emergent managed email proxy (constant — never read base URL from env)
EMAIL_BASE_URL = "https://integrations.emergentagent.com"
EMERGENT_EMAIL_KEY = os.environ.get('EMERGENT_EMAIL_KEY', '').strip()
EMAIL_FROM_NAME = os.environ.get('EMAIL_FROM_NAME', '1471 Horwich Squadron')

# Sending mailboxes on the verified domain (Resend). The local part is chosen
# per email type / sender role so recipients see who it's from.
EMAIL_DOMAIN = os.environ.get('EMAIL_DOMAIN', '1471squadron.co.uk')
APPOINTMENT_MAILBOX = {
    "training_officer": "trainingofficer",
    "adjutant": "adjutant",
    "stores_officer": "stores",
    "community_officer": "community",
    "health_safety_officer": "healthsafety",
    "shooting_officer": "shooting",
    "stem_officer": "stem",
    "oc": "oc",
    "deputy_oc": "deputyoc",
    "leadership_officer": "leadership",
    "sports_officer": "sports",
    "sqn_wo": "sqnwo",
    "dofe_officer": "dofe",
    "adventure_training_officer": "adventuretraining",
    "fieldcraft_officer": "fieldcraft",
    "cyber_officer": "cyber",
    "space_officer": "space",
}


def mailbox(local: str) -> str:
    return f"{local}@{EMAIL_DOMAIN}"

AGE_BANDS = {
    "under_12": "12 and under",
    "yr7_starting_yr8": "In Year 7, starting Year 8 in September",
    "yr8": "12 and in Year 8",
    "13_plus": "13 and over",
}

STAFF_ROLES = {"admin", "cfav"}
ALL_ROLES = {"admin", "cfav", "cadet", "parent"}

logging.basicConfig(level=logging.INFO,
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

app = FastAPI(title="1471 Horwich Squadron RAF Air Cadets API")
api_router = APIRouter(prefix="/api")
security = HTTPBearer(auto_error=False)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode("utf-8"), bcrypt.gensalt()).decode("utf-8")


def verify_password(plain: str, hashed: str) -> bool:
    return bcrypt.checkpw(plain.encode("utf-8"), hashed.encode("utf-8"))


def _clean_username_part(value: str) -> str:
    return re.sub(r"[^a-z0-9]", "", (value or "").lower())


def login_username_base(first_name: str, last_name: str) -> str:
    surname = _clean_username_part(last_name)
    first_initial = _clean_username_part(first_name)[:1]
    base = (surname + first_initial).strip()
    if not base:
        base = _clean_username_part(first_name) or "cadet"
    return base[:24]


async def ensure_login_username(first_name: str, last_name: str, exclude_user_id: Optional[str] = None) -> str:
    base = login_username_base(first_name, last_name)
    candidate = base
    suffix = 2
    while True:
        existing = await db.users.find_one({"login_username": candidate}, {"_id": 0, "id": 1})
        if not existing or existing.get("id") == exclude_user_id:
            return candidate
        candidate = f"{base}{suffix}"
        suffix += 1


def create_access_token(user_id: str, email: str, role: str) -> str:
    payload = {
        "sub": user_id, "email": email, "role": role,
        "exp": datetime.now(timezone.utc) + timedelta(hours=12), "type": "access",
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)


def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


async def get_current_user(creds: Optional[HTTPAuthorizationCredentials] = Depends(security)) -> dict:
    if creds is None:
        raise HTTPException(status_code=401, detail="Not authenticated")
    try:
        payload = jwt.decode(creds.credentials, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        if payload.get("type") != "access":
            raise HTTPException(status_code=401, detail="Invalid token type")
        user = await db.users.find_one({"id": payload["sub"]}, {"_id": 0, "password_hash": 0})
        if not user:
            raise HTTPException(status_code=401, detail="User not found")
        return user
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")


def require_roles(*roles):
    async def checker(user: dict = Depends(get_current_user)) -> dict:
        if user["role"] not in roles:
            raise HTTPException(status_code=403, detail="Insufficient permissions")
        return user
    return checker


require_staff = require_roles("admin", "cfav")


async def require_privileged_staff(user: dict = Depends(require_staff)) -> dict:
    if user.get("role") == "cfav" and not bool(user.get("is_uniformed", False)):
        raise HTTPException(status_code=403, detail="Uniformed CFAV or admin access required")
    return user


APPOINTMENT_KEYS = [
    "training_officer",
    "adjutant",
    "stores_officer",
    "community_officer",
    "health_safety_officer",
    "shooting_officer",
    "stem_officer",
    "oc",
    "deputy_oc",
    "leadership_officer",
    "sports_officer",
    "sqn_wo",
    "dofe_officer",
    "adventure_training_officer",
    "fieldcraft_officer",
    "cyber_officer",
    "space_officer",
]


async def _appointments_doc() -> dict:
    doc = await db.settings.find_one({"key": "appointments"}, {"_id": 0})
    return doc or {"key": "appointments", "value": {}}


async def _user_appointments(user_id: str) -> List[str]:
    app_doc = await _appointments_doc()
    v = app_doc.get("value", {})
    return [k for k in APPOINTMENT_KEYS if v.get(k) == user_id]


async def staff_from_email(staff: dict) -> str:
    """Pick the sending mailbox for a staff member: their appointment mailbox
    if they hold one, otherwise the site admin mailbox."""
    apps = await _user_appointments(staff.get("id", ""))
    for k in APPOINTMENT_KEYS:
        if k in apps and k in APPOINTMENT_MAILBOX:
            return mailbox(APPOINTMENT_MAILBOX[k])
    return mailbox("admin")


async def compute_member_stats(user_id: str) -> dict:
    """Points = attended-event points + bonus. Streak = attended volunteer events."""
    events = await db.events.find({"attendees": user_id}, {"_id": 0}).to_list(2000)
    points = sum(int(e.get("points_value", 0)) for e in events)
    streak = sum(1 for e in events if e.get("participation") == "volunteer")
    bonus = 0
    u = await db.users.find_one({"id": user_id}, {"_id": 0, "bonus_points": 1})
    if u:
        bonus = int(u.get("bonus_points", 0))
    return {"points": points + bonus, "event_points": points, "bonus_points": bonus,
            "streak": streak, "events_attended": len(events)}


def public_user(u: dict) -> dict:
    return {
        "id": u["id"], "email": u["email"], "role": u["role"],
        "login_username": u.get("login_username", ""),
        "must_change_password": bool(u.get("must_change_password", False)),
        "last_login_at": u.get("last_login_at"),
        "first_name": u.get("first_name", ""), "last_name": u.get("last_name", ""),
        "is_uniformed": bool(u.get("is_uniformed", False)) if u.get("role") == "cfav" else None,
        "child_ids": u.get("child_ids", []), "bonus_points": int(u.get("bonus_points", 0)),
        "dofe_level": u.get("dofe_level", "") if u.get("role") == "cadet" else "",
        "dofe_status": u.get("dofe_status", "") if u.get("role") == "cadet" else "",
        "btech_pathway": u.get("btech_pathway", "") if u.get("role") == "cadet" else "",
        "btech_status": u.get("btech_status", "") if u.get("role") == "cadet" else "",
        "cadet_notes": u.get("cadet_notes", "") if u.get("role") == "cadet" else "",
        "major_badges": u.get("major_badges", []) if u.get("role") == "cadet" else [],
        "created_at": u.get("created_at"),
    }


# ---------------------------------------------------------------------------
# Models
# ---------------------------------------------------------------------------
class LoginRequest(BaseModel):
    email: str
    password: str


class ChangePassword(BaseModel):
    current_password: str
    new_password: str = Field(..., min_length=6)


class ResetPassword(BaseModel):
    new_password: str = Field(..., min_length=6)


class UserCreate(BaseModel):
    email: Optional[EmailStr] = None
    first_name: str = Field(..., min_length=1)
    last_name: str = Field(default="")
    role: str
    is_uniformed: Optional[bool] = None
    password: str = Field(default="", min_length=0)
    child_ids: List[str] = []
    model_config = ConfigDict(extra="ignore")


class UserUpdate(BaseModel):
    first_name: Optional[str] = None
    last_name: Optional[str] = None
    role: Optional[str] = None
    is_uniformed: Optional[bool] = None
    child_ids: Optional[List[str]] = None
    bonus_points: Optional[int] = None
    dofe_level: Optional[str] = None
    dofe_status: Optional[str] = None
    btech_pathway: Optional[str] = None
    btech_status: Optional[str] = None
    cadet_notes: Optional[str] = None
    major_badges: Optional[List[str]] = None


class SelfRegisterRequest(BaseModel):
    role: str
    first_name: str = Field(..., min_length=1, max_length=120)
    last_name: str = Field(default="", max_length=120)
    email: Optional[EmailStr] = None
    is_uniformed: Optional[bool] = None
    model_config = ConfigDict(extra="ignore")


class CfavAvailabilityCreate(BaseModel):
    parade_date: str  # ISO date
    available: bool = True
    capabilities: List[str] = []
    note: str = Field(default="", max_length=2000)
    model_config = ConfigDict(extra="ignore")


class CfavEventIdeaCreate(BaseModel):
    title: str = Field(..., min_length=2, max_length=200)
    parade_date: Optional[str] = None
    summary: str = Field(default="", max_length=4000)
    model_config = ConfigDict(extra="ignore")


class CfavSkillMatrixCreate(BaseModel):
    skills: List[str] = []
    qualifications: List[str] = []
    interested_activities: List[str] = []
    willing_to_support: List[str] = []
    note: str = Field(default="", max_length=2000)
    model_config = ConfigDict(extra="ignore")


class TrainingPlanSlotUpdate(BaseModel):
    first_period_activity: str = Field(default="", max_length=300)
    second_period_activity: str = Field(default="", max_length=300)
    uniform_needed: str = Field(default="", max_length=200)
    model_config = ConfigDict(extra="ignore")


class TrainingPlanBidCreate(BaseModel):
    title: str = Field(..., min_length=2, max_length=200)
    summary: str = Field(default="", max_length=3000)
    model_config = ConfigDict(extra="ignore")


class TrainingPlanBidAccept(BaseModel):
    period: str = Field(default="first")  # first | second
    model_config = ConfigDict(extra="ignore")


class AppointmentsUpdate(BaseModel):
    training_officer: Optional[str] = None
    adjutant: Optional[str] = None
    stores_officer: Optional[str] = None
    community_officer: Optional[str] = None
    health_safety_officer: Optional[str] = None
    shooting_officer: Optional[str] = None
    stem_officer: Optional[str] = None
    oc: Optional[str] = None
    deputy_oc: Optional[str] = None
    leadership_officer: Optional[str] = None
    sports_officer: Optional[str] = None
    sqn_wo: Optional[str] = None
    dofe_officer: Optional[str] = None
    adventure_training_officer: Optional[str] = None
    fieldcraft_officer: Optional[str] = None
    cyber_officer: Optional[str] = None
    space_officer: Optional[str] = None
    model_config = ConfigDict(extra="ignore")


class EnquiryCreate(BaseModel):
    name: str = Field(..., min_length=2, max_length=120)
    email: EmailStr
    phone: Optional[str] = Field(default="", max_length=40)
    enquiry_type: str
    message: str = Field(..., min_length=5, max_length=4000)
    consent: bool
    dob: Optional[str] = None
    age_band: Optional[str] = None
    model_config = ConfigDict(extra="ignore")


class Enquiry(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    email: str
    phone: str = ""
    enquiry_type: str
    message: str
    consent: bool
    dob: Optional[str] = None
    age_band: Optional[str] = None
    status: str = "new"
    created_at: str = Field(default_factory=now_iso)


class StatusUpdate(BaseModel):
    status: str


class EventCreate(BaseModel):
    title: str = Field(..., min_length=2)
    description: str = ""
    location: str = ""
    link_url: str = ""
    attachment_ids: List[str] = []
    start: str  # ISO datetime
    end: Optional[str] = None
    capacity: int = 0  # 0 = unlimited
    event_type: str = "standard"  # standard | premium
    participation: str = "attend"  # attend | volunteer
    points_value: int = 10
    model_config = ConfigDict(extra="ignore")


class EventUpdate(BaseModel):
    title: Optional[str] = None
    description: Optional[str] = None
    location: Optional[str] = None
    link_url: Optional[str] = None
    attachment_ids: Optional[List[str]] = None
    start: Optional[str] = None
    end: Optional[str] = None
    capacity: Optional[int] = None
    event_type: Optional[str] = None
    participation: Optional[str] = None
    points_value: Optional[int] = None


class AttendanceUpdate(BaseModel):
    attendee_ids: List[str]


class NoticeCreate(BaseModel):
    title: str = Field(..., min_length=2)
    body: str = Field(..., min_length=1)
    roles: List[str]
    requires_ack: bool = False
    model_config = ConfigDict(extra="ignore")


class MessageCreate(BaseModel):
    body: str = Field(..., min_length=1, max_length=4000)


class SmsResetRequestCreate(BaseModel):
    note: str = Field(default="", max_length=2000)
    model_config = ConfigDict(extra="ignore")


class Audience(BaseModel):
    mode: str = "all"  # all | roles | users | parent_of
    roles: List[str] = []
    user_ids: List[str] = []
    cadet_id: Optional[str] = None
    appointment_keys: List[str] = []
    model_config = ConfigDict(extra="ignore")


class BroadcastCreate(BaseModel):
    title: str = Field(..., min_length=2, max_length=200)
    body: str = Field(..., min_length=1, max_length=8000)
    audience: Audience = Field(default_factory=Audience)
    channels: List[str] = ["dashboard"]
    attachment_ids: List[str] = []
    base_url: str = ""
    model_config = ConfigDict(extra="ignore")


class NewsletterCreate(BaseModel):
    subject: str = Field(..., min_length=2, max_length=200)
    heading: str = Field(default="", max_length=200)
    intro: str = Field(default="", max_length=4000)
    body: str = Field(..., min_length=1, max_length=20000)
    attachment_ids: List[str] = []
    model_config = ConfigDict(extra="ignore")


class NewsletterSend(BaseModel):
    audience: Audience = Field(default_factory=Audience)
    channels: List[str] = ["dashboard", "email"]
    base_url: str = ""
    model_config = ConfigDict(extra="ignore")


class SiteImageOverride(BaseModel):
    src: str = Field(default="", max_length=4000)
    alt: str = Field(default="", max_length=1000)
    model_config = ConfigDict(extra="ignore")


class SiteContentUpdate(BaseModel):
    path: str = Field(..., min_length=1, max_length=200)
    texts: Dict[str, str] = {}
    images: Dict[str, SiteImageOverride] = {}
    model_config = ConfigDict(extra="ignore")


# ---------------------------------------------------------------------------
# Email (Emergent managed email proxy)
# ---------------------------------------------------------------------------
def _email_shell(title: str, inner: str) -> str:
    return f"""
    <table width="100%" cellpadding="0" cellspacing="0" style="font-family:Arial,Helvetica,sans-serif;background:#EAF5F8;padding:24px;">
      <tr><td align="center"><table width="600" style="background:#fff;border:1px solid #d6e6ec;">
        <tr><td style="background:#002F5F;padding:20px 28px;color:#fff;font-size:18px;font-weight:bold;">1471 Horwich Squadron RAF Air Cadets</td></tr>
        <tr><td style="height:4px;background:#C60C30;"></td></tr>
        <tr><td style="padding:28px;color:#071A2F;font-size:15px;line-height:1.6;">
          <h2 style="margin:0 0 14px;color:#002F5F;font-size:20px;">{title}</h2>
          {inner}
        </td></tr>
        <tr><td style="padding:16px 28px;background:#F3F8FA;color:#5b6b78;font-size:12px;">1471 Horwich Squadron RAF Air Cadets &middot; This message was sent from the Squadron members area.</td></tr>
      </table></td></tr></table>"""


def _enquiry_email_html(e: Enquiry) -> str:
    band = AGE_BANDS.get(e.age_band or "", "")
    extra = ""
    if e.dob or band:
        extra = f"""<p><strong>Date of birth:</strong> {e.dob or '&ndash;'}</p>
          <p><strong>Age band:</strong> {band or '&ndash;'}</p>"""
    inner = f"""
      <p><strong>Type:</strong> {e.enquiry_type}</p>
      <p><strong>Name:</strong> {e.name}</p>
      <p><strong>Email:</strong> {e.email}</p>
      <p><strong>Phone:</strong> {e.phone or '&ndash;'}</p>
      {extra}
      <p><strong>Message:</strong></p>
      <p style="padding:14px;background:#EAF5F8;border-left:3px solid #002F5F;">{e.message}</p>"""
    return _email_shell("New website enquiry", inner)


def _text_to_html(text: str) -> str:
    safe = (text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    paras = [p.strip() for p in safe.split("\n\n") if p.strip()]
    return "".join(f'<p style="margin:0 0 12px;">{p.replace(chr(10), "<br/>")}</p>' for p in paras)


def _attachments_email_html(attachments: Optional[list], base_url: str = "") -> str:
    if not attachments:
        return ""
    items = "".join(
        f'<li style="margin:4px 0;"><a href="{base_url}/api/attachments/{a["id"]}/download" style="color:#002F5F;">{a["filename"]}</a></li>'
        for a in attachments)
    return f'<p><strong>Attachments</strong></p><ul style="padding-left:20px;">{items}</ul>'


def _broadcast_email_html(title: str, body: str, from_name: str,
                          attachments: Optional[list] = None, base_url: str = "") -> str:
    inner = (
        _text_to_html(body)
        + _attachments_email_html(attachments, base_url)
        + f'<p style="margin-top:20px;color:#5b6b78;font-size:13px;">&mdash; {from_name}</p>'
    )
    return _email_shell(title, inner)


def _newsletter_email_html(nl: dict, attachments: Optional[list] = None, base_url: str = "") -> str:
    parts = []
    if nl.get("intro"):
        parts.append(f'<p style="font-size:16px;color:#334;">{_text_to_html(nl["intro"])}</p>')
    parts.append(_text_to_html(nl.get("body", "")))
    parts.append(_attachments_email_html(attachments, base_url))
    return _email_shell(nl.get("heading") or nl.get("subject", "Squadron newsletter"), "".join(parts))


async def send_email(to: str, subject: str, html: str, reply_to: Optional[str] = None,
                     from_email: Optional[str] = None, from_name: Optional[str] = None) -> str:
    """Send one email. Prefers Resend (own domain); falls back to the managed proxy.
    `from_email` sets the sending mailbox; defaults to SENDER_EMAIL.
    Returns 'sent' | 'skipped' | 'failed'."""
    sender = from_email or SENDER_EMAIL
    name = from_name or EMAIL_FROM_NAME
    # Preferred: Resend (user's own verified domain)
    if RESEND_API_KEY and resend is not None:
        resend.api_key = RESEND_API_KEY
        params = {
            "from": f"{name} <{sender}>",
            "to": [to],
            "subject": subject,
            "html": html,
        }
        if reply_to:
            params["reply_to"] = reply_to
        try:
            await asyncio.to_thread(resend.Emails.send, params)
            return "sent"
        except Exception as exc:  # pragma: no cover
            logger.error("Resend send failed to %s: %s", to, exc)
            return "failed"
    # Fallback: Emergent managed email proxy
    if not EMERGENT_EMAIL_KEY:
        logger.info("Email not configured; skipped send to %s", to)
        return "skipped"
    payload = {"to": [to], "subject": subject, "html": html, "from_name": name}
    if reply_to:
        payload["contact_email"] = reply_to
    try:
        async with httpx.AsyncClient(timeout=30) as c:
            resp = await c.post(f"{EMAIL_BASE_URL}/api/v1/email/send",
                                headers={"X-Email-Key": EMERGENT_EMAIL_KEY}, json=payload)
        resp.raise_for_status()
        return "sent"
    except Exception as exc:  # pragma: no cover
        logger.error("Email send failed to %s: %s", to, exc)
        return "failed"


async def send_enquiry_email(e: Enquiry) -> None:
    await send_email(NOTIFY_EMAIL, f"New {e.enquiry_type} enquiry - {e.name}",
                     _enquiry_email_html(e), reply_to=e.email, from_email=mailbox("admin"))


# ---------------------------------------------------------------------------
# Web Push (VAPID) — notifications + app icon badge on installed PWA
# ---------------------------------------------------------------------------
from pywebpush import webpush, WebPushException  # noqa: E402

VAPID_PUBLIC_KEY = os.environ.get("VAPID_PUBLIC_KEY", "").strip()
VAPID_SUBJECT = os.environ.get("VAPID_SUBJECT", "mailto:admin@1471horwich.org.uk")


def _load_vapid_pem() -> Optional[str]:
    b64 = os.environ.get("VAPID_PRIVATE_PEM_B64", "").strip()
    if not b64:
        return None
    path = "/tmp/vapid_private.pem"
    with open(path, "wb") as f:
        f.write(base64.b64decode(b64))
    return path


VAPID_PEM_PATH = _load_vapid_pem()


def _send_one_push(subscription: dict, payload: str) -> None:
    webpush(subscription_info=subscription, data=payload,
            vapid_private_key=VAPID_PEM_PATH, vapid_claims={"sub": VAPID_SUBJECT})


async def push_to_user(user_id: str, title: str, body: str, url: str = "/portal") -> None:
    if not VAPID_PEM_PATH:
        return
    subs = await db.push_subscriptions.find({"user_id": user_id}).to_list(20)
    if not subs:
        return
    badge = await db.notifications.count_documents({"user_id": user_id, "read": False})
    payload = json.dumps({"title": title, "body": (body or "")[:180], "url": url, "badge": badge})
    for s in subs:
        try:
            await asyncio.to_thread(_send_one_push, s["subscription"], payload)
        except WebPushException as exc:
            code = getattr(getattr(exc, "response", None), "status_code", None)
            if code in (404, 410):
                await db.push_subscriptions.delete_one({"_id": s["_id"]})
        except Exception as exc:  # pragma: no cover
            logger.error("Push failed for %s: %s", user_id, exc)


# ---------------------------------------------------------------------------
# Recruitment eligibility + broadcast helpers
# ---------------------------------------------------------------------------
def _eligible_date_obj(e: dict):
    """Canonical date a prospective cadet becomes eligible (date obj), or None."""
    band = e.get("age_band")
    if not band:
        return None
    try:
        cd = datetime.fromisoformat(e.get("created_at", "")).date()
    except Exception:
        cd = date.today()
    if band in ("yr8", "13_plus"):
        return cd  # already eligible when they enquired
    if band == "yr7_starting_yr8":
        s = date(cd.year, 9, 1)
        if cd > s:
            s = date(cd.year + 1, 9, 1)
        return s
    if band == "under_12":
        dob = e.get("dob")
        if not dob:
            return None
        try:
            d = datetime.fromisoformat(dob).date()
        except Exception:
            try:
                d = datetime.strptime(dob, "%Y-%m-%d").date()
            except Exception:
                return None
        # England school-year rule: Year 8 starts the September of birth_year + (12 if born Sep-Dec else 11)
        yr = d.year + (12 if d.month >= 9 else 11)
        return date(yr, 9, 1)
    return None


def eligible_date(e: dict) -> Optional[str]:
    ed = _eligible_date_obj(e)
    return ed.isoformat() if ed else None


def compute_eligibility(e: dict) -> Optional[str]:
    """Time-aware bucket: prospects move now <- september <- future as time passes."""
    if not e.get("age_band"):
        return None
    ed = _eligible_date_obj(e)
    if ed is None:
        return "future"  # cadet enquiry with no date of birth on record
    today = date.today()
    if ed <= today:
        return "now"
    ns = date(today.year, 9, 1)
    if today > ns:
        ns = date(today.year + 1, 9, 1)
    return "september" if ed <= ns else "future"


def _dob_age_years(dob_value: Optional[str]) -> Optional[int]:
    """Age in full years from ISO dob; returns None if missing/invalid."""
    if not dob_value:
        return None
    try:
        d = datetime.fromisoformat(dob_value).date()
    except Exception:
        try:
            d = datetime.strptime(dob_value, "%Y-%m-%d").date()
        except Exception:
            return None
    today = date.today()
    years = today.year - d.year - ((today.month, today.day) < (d.month, d.day))
    return years


def age_mismatch_info(e: dict) -> dict:
    """Detect obvious mismatch between selected age band and supplied DoB.

    This flags clear contradictions only (for follow-up), not borderline school-year cases.
    """
    band = e.get("age_band")
    age = _dob_age_years(e.get("dob"))
    if not band or age is None:
        return {"is_mismatch": False, "reason": "", "expected": ""}

    expected = ""
    mismatch = False
    reason = ""

    if age >= 13:
        expected = "13_plus"
        if band != "13_plus":
            mismatch = True
            reason = "Selected band is younger than DoB-based age."
    elif age == 12:
        expected = "yr8"
        if band in ("13_plus", "under_12"):
            mismatch = True
            reason = "Selected band does not match DoB-based age of 12."
    else:
        expected = "under_12"
        if band in ("yr8", "13_plus"):
            mismatch = True
            reason = "Selected band is older than DoB-based age."

    return {"is_mismatch": mismatch, "reason": reason, "expected": expected}


def countdown_text(target_iso: str) -> str:
    target = date.fromisoformat(target_iso)
    days = (target - date.today()).days
    if days <= 0:
        return "you can join us now"
    months, rem = days // 30, days % 30
    if months >= 1:
        out = f"about {months} month{'s' if months != 1 else ''}"
        if rem:
            out += f", {rem} day{'s' if rem != 1 else ''}"
        return out
    return f"{days} day{'s' if days != 1 else ''}"


def _joining_instructions_html(name: str, note: str = "", attachments: Optional[list] = None, base_url: str = "") -> str:
    note_html = f'<p style="padding:12px;background:#EAF5F8;border-left:3px solid #002F5F;">{note}</p>' if note else ""
    docs_html = ""
    if attachments:
        items = "".join(
            f'<li style="margin:4px 0;"><a href="{base_url}/api/attachments/{a["id"]}/download" style="color:#002F5F;">{a["filename"]}</a></li>'
            for a in attachments)
        docs_html = f'<p><strong>Documents to download</strong></p><ul style="padding-left:20px;">{items}</ul>'
    inner = f"""
      <p>Hi {name},</p>
      <p>Great news &mdash; you&rsquo;re old enough to join 1471 Horwich Squadron RAF Air Cadets, and we&rsquo;d love to welcome you along.</p>
      {note_html}
      <p><strong>When we parade</strong><br/>Monday &amp; Thursday evenings, 19:00&ndash;21:30.</p>
      <p><strong>Where</strong><br/>1471 (Horwich) ATC Sqn HQ, St Joseph&rsquo;s Secondary School &amp; Sports College, Chorley New Road, Horwich, BL6 6HW.</p>
      {docs_html}
      <p><strong>What to do next</strong><br/>Come along on a parade night with a parent or carer, in comfortable clothing. You don&rsquo;t need any experience or kit. Please reply to this email to let us know which evening you&rsquo;ll visit so we can look out for you.</p>
      <p>We look forward to meeting you!</p>
      <p>1471 Horwich Squadron staff team</p>"""
    return _email_shell("You can join us &mdash; here&rsquo;s how", inner)


def _countdown_html(name: str, target_iso: str, note: str = "") -> str:
    when = date.fromisoformat(target_iso).strftime("%d %B %Y").lstrip("0")
    note_html = f'<p style="padding:12px;background:#EAF5F8;border-left:3px solid #002F5F;">{note}</p>' if note else ""
    inner = f"""
      <p>Hi {name},</p>
      <p>Thanks for your interest in joining 1471 Horwich Squadron RAF Air Cadets!</p>
      <p>You&rsquo;re not quite old enough to join just yet &mdash; cadets can normally start from Year 8 (age 12). Based on the details you gave us, you&rsquo;ll be able to join us from <strong>{when}</strong>.</p>
      <p style="font-size:18px;color:#C60C30;font-weight:bold;">That&rsquo;s {countdown_text(target_iso)} to go!</p>
      {note_html}
      <p>We&rsquo;ll be in touch nearer the time. In the meantime, follow us on Facebook to see what our cadets get up to.</p>
      <p>See you soon,<br/>1471 Horwich Squadron staff team</p>"""
    return _email_shell("Not long to go until you can join!", inner)


async def send_recruit_email(e: dict, kind: Optional[str], note: str,
                             attachments: Optional[list] = None, base_url: str = "") -> dict:
    first = (e.get("name", "there").strip().split(" ") or ["there"])[0]
    elig = compute_eligibility(e)
    k = "joining" if elig == "now" else (kind if kind in ("joining", "countdown") else "countdown")
    if k == "joining":
        status = await send_email(e["email"], "Joining 1471 Horwich Squadron \u2014 here's how",
                                  _joining_instructions_html(first, note, attachments, base_url),
                                  reply_to=mailbox("enrolments"), from_email=mailbox("enrolments"))
        target = None
    else:
        target = eligible_date(e)
        if not target:
            return {"sent": False, "error": "no_date"}
        status = await send_email(e["email"], "Not long until you can join 1471 Horwich Squadron!",
                                  _countdown_html(first, target, note),
                                  reply_to=mailbox("enrolments"), from_email=mailbox("enrolments"))
    set_fields = {"last_recruit_email": {"kind": k, "at": now_iso(), "status": status}}
    if status == "sent":
        set_fields["status"] = "actioned"
    await db.enquiries.update_one({"id": e["id"]}, {"$set": set_fields})
    return {"sent": status == "sent", "email_status": status, "kind": k, "eligible_date": target}


async def resolve_recipients(a: Audience) -> List[dict]:
    if a.mode == "all":
        q = {}
    elif a.mode == "roles":
        roles = [r for r in a.roles if r in ALL_ROLES]
        if not roles:
            return []
        q = {"role": {"$in": roles}}
    elif a.mode == "users":
        if not a.user_ids:
            return []
        q = {"id": {"$in": a.user_ids}}
    elif a.mode == "parent_of":
        if not a.cadet_id:
            return []
        q = {"role": "parent", "child_ids": a.cadet_id}
    elif a.mode == "appointments":
        keys = [k for k in a.appointment_keys if k in APPOINTMENT_KEYS]
        if not keys:
            return []
        app_doc = await _appointments_doc()
        value = app_doc.get("value", {})
        user_ids = sorted({value.get(k) for k in keys if value.get(k)})
        if not user_ids:
            return []
        q = {"id": {"$in": user_ids}}
    else:
        return []
    return await db.users.find(q, {"_id": 0, "password_hash": 0}).to_list(5000)


async def deliver_broadcast(users: List[dict], title: str, body: str, channels: List[str],
                            from_name: str, kind: str, email_html: Optional[str] = None,
                            link: Optional[str] = None, link_label: Optional[str] = None,
                            links: Optional[List[dict]] = None, from_email: Optional[str] = None) -> dict:
    channels = [c for c in channels if c in ("dashboard", "email")] or ["dashboard"]
    docs, emails_sent = [], 0
    for u in users:
        email_status = "n/a"
        if "email" in channels and u.get("email"):
            html = email_html or _broadcast_email_html(title, body, from_name)
            email_status = await send_email(u["email"], title, html, from_email=from_email)
            if email_status == "sent":
                emails_sent += 1
        if "dashboard" in channels:
            docs.append({"id": str(uuid.uuid4()), "user_id": u["id"], "title": title,
                         "body": body, "from_name": from_name, "kind": kind,
                         "channels": channels, "email_status": email_status,
                         "link": link, "link_label": link_label,
                         "links": links or [],
                         "read": False, "created_at": now_iso()})
    if docs:
        await db.notifications.insert_many(docs)
        for d in docs:
            await push_to_user(d["user_id"], title, d.get("body", ""), "/portal")
    return {"recipients": len(users), "dashboard_delivered": len(docs), "emails_sent": emails_sent,
            "email_configured": bool(RESEND_API_KEY or EMERGENT_EMAIL_KEY)}


# ---------------------------------------------------------------------------
# Public + Auth routes
# ---------------------------------------------------------------------------
@api_router.get("/")
async def root():
    return {"message": "1471 Horwich Squadron RAF Air Cadets API"}


@api_router.post("/auth/login")
async def login(req: LoginRequest):
    ident = (req.email or "").strip().lower()
    if not ident:
        raise HTTPException(status_code=400, detail="Email or username is required")
    user = await db.users.find_one({"$or": [{"email": ident}, {"login_username": ident}]})
    if not user or not verify_password(req.password, user["password_hash"]):
        raise HTTPException(status_code=401, detail="Invalid username/email or password")
    await db.users.update_one(
        {"id": user["id"]},
        {"$set": {"last_login_at": now_iso(), "login_reminder_sent_at": None}},
    )
    user["last_login_at"] = now_iso()
    token = create_access_token(user["id"], user["email"], user["role"])
    return {"access_token": token, "token_type": "bearer", "user": public_user(user)}


@api_router.get("/auth/me")
async def me(user: dict = Depends(get_current_user)):
    if not user.get("login_username"):
        generated = await ensure_login_username(user.get("first_name", ""), user.get("last_name", ""), user.get("id"))
        await db.users.update_one({"id": user["id"]}, {"$set": {"login_username": generated}})
        user["login_username"] = generated
    data = public_user(user)
    if user["role"] in ("cadet", "parent", "cfav", "admin"):
        data["stats"] = await compute_member_stats(user["id"])
    data["appointments"] = await _user_appointments(user["id"])
    if user["role"] == "cadet":
        data["initial_password"] = CADET_DEFAULT_PASSWORD
    return data


async def _event_detail_view(e: dict, user_id: str, staff: bool) -> dict:
    view = event_view(e, user_id, staff)
    view["attachments"] = await _fetch_attachments(e.get("attachment_ids", []))
    if staff and e.get("bids"):
        users = await db.users.find({"id": {"$in": e.get("bids", [])}}, {"_id": 0, "id": 1, "first_name": 1, "last_name": 1, "role": 1, "is_uniformed": 1}).to_list(500)
        by_id = {u["id"]: u for u in users}
        view["bidders"] = [by_id[bid] for bid in e.get("bids", []) if bid in by_id]
    else:
        view["bidders"] = []
    return view


@api_router.post("/auth/change-password")
async def change_password(payload: ChangePassword, user: dict = Depends(get_current_user)):
    full = await db.users.find_one({"id": user["id"]})
    if not verify_password(payload.current_password, full["password_hash"]):
        raise HTTPException(status_code=400, detail="Current password is incorrect")
    await db.users.update_one({"id": user["id"]},
                              {"$set": {"password_hash": hash_password(payload.new_password),
                                        "must_change_password": False}})
    return {"updated": True}


# ---------------------------------------------------------------------------
# Enquiries
# ---------------------------------------------------------------------------
@api_router.post("/enquiries", response_model=Enquiry)
async def create_enquiry(payload: EnquiryCreate):
    if not payload.consent:
        raise HTTPException(status_code=400, detail="Consent is required to submit this form.")
    enquiry = Enquiry(**payload.model_dump())
    await db.enquiries.insert_one(enquiry.model_dump())
    await send_enquiry_email(enquiry)
    return enquiry


def _enquiry_out(e: dict) -> dict:
    e = {k: v for k, v in e.items() if k != "_id"}
    e["eligibility"] = compute_eligibility(e)
    e["age_band_label"] = AGE_BANDS.get(e.get("age_band") or "", "")
    e["eligible_date"] = eligible_date(e)
    mm = age_mismatch_info(e)
    e["age_mismatch"] = bool(mm["is_mismatch"])
    e["age_mismatch_reason"] = mm["reason"]
    e["expected_age_band"] = mm["expected"]
    e["expected_age_band_label"] = AGE_BANDS.get(mm["expected"] or "", "")
    return e


@api_router.get("/enquiries")
async def list_enquiries(staff: dict = Depends(require_privileged_staff)):
    rows = await db.enquiries.find({}, {"_id": 0}).sort("created_at", -1).to_list(1000)
    return [_enquiry_out(e) for e in rows]


@api_router.get("/enquiries/tracker")
async def enquiries_tracker(staff: dict = Depends(require_privileged_staff)):
    """Prospective-cadet enquiries grouped by joining eligibility."""
    rows = await db.enquiries.find({"age_band": {"$ne": None}}, {"_id": 0}).sort("created_at", -1).to_list(2000)
    buckets = {"now": [], "september": [], "future": []}
    follow_up_age_mismatch = []
    for e in rows:
        out = _enquiry_out(e)
        if out["eligibility"] in buckets:
            buckets[out["eligibility"]].append(out)
        if out.get("age_mismatch"):
            follow_up_age_mismatch.append(out)
    return {"buckets": buckets,
            "counts": {k: len(v) for k, v in buckets.items()},
            "labels": {"now": "Can join now", "september": "Eligible in September",
                       "future": "Eligible in the future"},
            "follow_up": {"age_mismatch": follow_up_age_mismatch},
            "follow_up_counts": {"age_mismatch": len(follow_up_age_mismatch)},
            "follow_up_labels": {"age_mismatch": "Follow Up - Age Mismatch"}}


@api_router.patch("/enquiries/{enquiry_id}", response_model=Enquiry)
async def update_enquiry(enquiry_id: str, update: StatusUpdate, staff: dict = Depends(require_privileged_staff)):
    if update.status not in {"new", "read", "actioned"}:
        raise HTTPException(status_code=400, detail="Invalid status")
    res = await db.enquiries.find_one_and_update(
        {"id": enquiry_id}, {"$set": {"status": update.status}},
        projection={"_id": 0}, return_document=True)
    if not res:
        raise HTTPException(status_code=404, detail="Enquiry not found")
    return res


@api_router.delete("/enquiries/{enquiry_id}")
async def delete_enquiry(enquiry_id: str, staff: dict = Depends(require_privileged_staff)):
    res = await db.enquiries.delete_one({"id": enquiry_id})
    if res.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Enquiry not found")
    return {"deleted": True}


class RecruitEmail(BaseModel):
    kind: Optional[str] = None  # joining | countdown
    note: str = Field(default="", max_length=2000)
    attachment_ids: List[str] = []
    base_url: str = ""
    model_config = ConfigDict(extra="ignore")


class RecruitBulk(BaseModel):
    eligibility: str
    note: str = Field(default="", max_length=2000)
    attachment_ids: List[str] = []
    base_url: str = ""
    model_config = ConfigDict(extra="ignore")


async def _fetch_attachments(ids: List[str]) -> list:
    if not ids:
        return []
    rows = await db.attachments.find({"id": {"$in": ids}}, {"_id": 0}).to_list(50)
    order = {aid: i for i, aid in enumerate(ids)}
    rows.sort(key=lambda a: order.get(a["id"], 999))
    return rows


@api_router.post("/enquiries/recruit-email/bulk")
async def recruit_email_bulk(payload: RecruitBulk, staff: dict = Depends(require_privileged_staff)):
    if payload.eligibility not in ("now", "september", "future"):
        raise HTTPException(status_code=400, detail="Invalid eligibility group")
    atts = await _fetch_attachments(payload.attachment_ids)
    rows = await db.enquiries.find({"age_band": {"$ne": None}}, {"_id": 0}).to_list(2000)
    sent, skipped = 0, 0
    for e in rows:
        if compute_eligibility(e) == payload.eligibility:
            res = await send_recruit_email(e, None, payload.note, atts, payload.base_url)
            if res.get("email_status") == "sent":
                sent += 1
            else:
                skipped += 1
    return {"sent": sent, "skipped": skipped,
            "kind": "joining" if payload.eligibility == "now" else "countdown"}


@api_router.post("/enquiries/{enquiry_id}/recruit-email")
async def recruit_email(enquiry_id: str, payload: RecruitEmail, staff: dict = Depends(require_privileged_staff)):
    e = await db.enquiries.find_one({"id": enquiry_id}, {"_id": 0})
    if not e:
        raise HTTPException(status_code=404, detail="Enquiry not found")
    atts = await _fetch_attachments(payload.attachment_ids)
    res = await send_recruit_email(e, payload.kind, payload.note, atts, payload.base_url)
    if not res.get("sent") and res.get("error") == "no_date":
        raise HTTPException(status_code=400,
                            detail="Cannot work out an eligibility date for this enquiry (no date of birth on record).")
    return res


# ---------------------------------------------------------------------------
# Attachments (GridFS-backed; download link is public/unguessable)
# ---------------------------------------------------------------------------
@api_router.post("/attachments")
async def upload_attachment(file: UploadFile = File(...), staff: dict = Depends(require_staff)):
    data = await file.read()
    if len(data) > 15 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="File too large (maximum 15MB).")
    aid = str(uuid.uuid4())
    gid = await fs.upload_from_stream(file.filename or "file", data,
                                      metadata={"attachment_id": aid, "content_type": file.content_type})
    doc = {"id": aid, "filename": file.filename or "file",
           "content_type": file.content_type or "application/octet-stream",
           "size": len(data), "gridfs_id": str(gid), "uploaded_by": staff["id"], "created_at": now_iso()}
    await db.attachments.insert_one(doc)
    return {k: v for k, v in doc.items() if k != "_id"}


@api_router.get("/attachments")
async def list_attachments(staff: dict = Depends(require_staff)):
    return await db.attachments.find({}, {"_id": 0}).sort("created_at", -1).to_list(500)


@api_router.get("/attachments/{attachment_id}/download")
async def download_attachment(attachment_id: str):
    doc = await db.attachments.find_one({"id": attachment_id}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Attachment not found")
    stream = await fs.open_download_stream(ObjectId(doc["gridfs_id"]))
    data = await stream.read()
    return Response(content=data, media_type=doc["content_type"],
                    headers={"Content-Disposition": f'inline; filename="{doc["filename"]}"'})


@api_router.delete("/attachments/{attachment_id}")
async def delete_attachment(attachment_id: str, staff: dict = Depends(require_staff)):
    doc = await db.attachments.find_one({"id": attachment_id})
    if not doc:
        raise HTTPException(status_code=404, detail="Attachment not found")
    try:
        await fs.delete(ObjectId(doc["gridfs_id"]))
    except Exception:
        pass
    await db.attachments.delete_one({"id": attachment_id})
    return {"deleted": True}


# ---------------------------------------------------------------------------
# Document library (GridFS-backed, browsable + shareable)
# ---------------------------------------------------------------------------
class DocumentSend(BaseModel):
    audience: Audience = Field(default_factory=Audience)
    channels: List[str] = ["dashboard", "email"]
    message: str = Field(default="", max_length=2000)
    base_url: str = ""
    model_config = ConfigDict(extra="ignore")


class DocumentMeta(BaseModel):
    title: str = Field(..., min_length=1, max_length=200)
    category: str = Field(default="General", max_length=80)
    visible_roles: List[str] = []
    model_config = ConfigDict(extra="ignore")


class LearningAssignmentCreate(BaseModel):
    document_id: str
    cadet_ids: List[str] = []
    instructions: str = Field(default="", max_length=4000)
    due_date: str = Field(default="", max_length=40)
    model_config = ConfigDict(extra="ignore")


def _document_out(d: dict) -> dict:
    return {k: v for k, v in d.items() if k not in ("_id", "gridfs_id")}


def _document_email_html(title: str, message: str, link: str, from_name: str) -> str:
    msg = f'<p>{message}</p>' if message else ""
    inner = f"""
      {msg}
      <p>A document has been shared with you: <strong>{title}</strong></p>
      <p><a href="{link}" style="display:inline-block;background:#C60C30;color:#fff;padding:12px 22px;text-decoration:none;font-weight:bold;">Download document</a></p>
      <p style="font-size:12px;color:#5b6b78;">Or copy this link: {link}</p>
      <p style="margin-top:18px;">&mdash; {from_name}</p>"""
    return _email_shell(title, inner)


async def _learning_assignment_out(a: dict, viewer: Optional[dict] = None) -> dict:
    out = {k: v for k, v in a.items() if k != "_id"}
    doc = await db.documents.find_one({"id": a.get("document_id")}, {"_id": 0, "gridfs_id": 0})
    out["document"] = _document_out(doc) if doc else None

    cadet_ids = a.get("cadet_ids", []) or []
    cadets = await db.users.find({"id": {"$in": cadet_ids}}, {"_id": 0, "id": 1, "first_name": 1, "last_name": 1}).to_list(500)
    by_cadet = {c["id"]: c for c in cadets}
    out["cadets"] = [{
        "id": cid,
        "first_name": by_cadet.get(cid, {}).get("first_name", ""),
        "last_name": by_cadet.get(cid, {}).get("last_name", ""),
    } for cid in cadet_ids]

    submissions = await db.learning_submissions.find({"assignment_id": a["id"]}, {"_id": 0, "gridfs_id": 0}).sort("submitted_at", -1).to_list(500)
    out["submitted_count"] = len(submissions)
    if viewer and viewer.get("role") == "cadet":
        out["my_submission"] = next((s for s in submissions if s.get("cadet_id") == viewer.get("id")), None)
    else:
        for s in submissions:
            cadet = by_cadet.get(s.get("cadet_id"), {})
            s["cadet_name"] = f"{cadet.get('first_name', '')} {cadet.get('last_name', '')}".strip()
        out["submissions"] = submissions
    return out


async def _learning_reviewer_ids() -> List[str]:
    app_doc = await _appointments_doc()
    reviewer_ids = []
    training_id = (app_doc.get("value") or {}).get("training_officer")
    if training_id:
        reviewer_ids.append(training_id)
    if not reviewer_ids:
        admins = await db.users.find({"role": "admin"}, {"_id": 0, "id": 1}).to_list(50)
        reviewer_ids.extend(a["id"] for a in admins)
    return list(dict.fromkeys(reviewer_ids))


@api_router.post("/documents")
async def upload_document(file: UploadFile = File(...), title: str = Form(...),
                          category: str = Form("General"), visible_roles: str = Form(""),
                          staff: dict = Depends(require_staff)):
    data = await file.read()
    if len(data) > 15 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="File too large (maximum 15MB).")
    roles = [r for r in visible_roles.split(",") if r in ALL_ROLES]
    gid = await fs.upload_from_stream(file.filename or "file", data,
                                      metadata={"content_type": file.content_type})
    doc = {"id": str(uuid.uuid4()), "title": title.strip(), "category": (category or "General").strip(),
           "filename": file.filename or "file", "content_type": file.content_type or "application/octet-stream",
           "size": len(data), "gridfs_id": str(gid), "visible_roles": roles,
           "uploaded_by": staff["id"],
           "uploaded_by_name": f"{staff.get('first_name','')} {staff.get('last_name','')}".strip() or "Staff",
           "created_at": now_iso()}
    await db.documents.insert_one(doc)
    return _document_out(doc)


@api_router.get("/documents")
async def list_documents(staff: dict = Depends(require_staff)):
    rows = await db.documents.find({}, {"gridfs_id": 0}).sort("created_at", -1).to_list(1000)
    return [_document_out(d) for d in rows]


@api_router.get("/documents/library")
async def library_documents(user: dict = Depends(get_current_user)):
    if user["role"] == "parent":
        roles = ["parent", "cadet"]
        rows = await db.documents.find({"visible_roles": {"$in": roles}}, {"gridfs_id": 0}).sort("created_at", -1).to_list(1000)
    else:
        rows = await db.documents.find({"visible_roles": user["role"]}, {"gridfs_id": 0}).sort("created_at", -1).to_list(1000)
    return [_document_out(d) for d in rows]


@api_router.patch("/documents/{document_id}")
async def update_document(document_id: str, payload: DocumentMeta, staff: dict = Depends(require_staff)):
    roles = [r for r in payload.visible_roles if r in ALL_ROLES]
    d = await db.documents.find_one_and_update(
        {"id": document_id},
        {"$set": {"title": payload.title.strip(), "category": (payload.category or "General").strip(),
                  "visible_roles": roles}},
        projection={"_id": 0, "gridfs_id": 0}, return_document=True)
    if not d:
        raise HTTPException(status_code=404, detail="Document not found")
    return d


@api_router.get("/documents/{document_id}/download")
async def download_document(document_id: str):
    doc = await db.documents.find_one({"id": document_id}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    stream = await fs.open_download_stream(ObjectId(doc["gridfs_id"]))
    data = await stream.read()
    return Response(content=data, media_type=doc["content_type"],
                    headers={"Content-Disposition": f'inline; filename="{doc["filename"]}"'})


@api_router.delete("/documents/{document_id}")
async def delete_document(document_id: str, staff: dict = Depends(require_staff)):
    doc = await db.documents.find_one({"id": document_id})
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    try:
        await fs.delete(ObjectId(doc["gridfs_id"]))
    except Exception:
        pass
    await db.documents.delete_one({"id": document_id})
    return {"deleted": True}


@api_router.post("/documents/{document_id}/send")
async def send_document(document_id: str, payload: DocumentSend, staff: dict = Depends(require_staff)):
    doc = await db.documents.find_one({"id": document_id}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    users = await resolve_recipients(payload.audience)
    if not users:
        raise HTTPException(status_code=400, detail="No recipients match this audience")
    from_name = f"{staff.get('first_name','')} {staff.get('last_name','')}".strip() or "Squadron Staff"
    from_email = await staff_from_email(staff)
    abs_link = f"{payload.base_url}/api/documents/{document_id}/download"
    rel_link = f"/api/documents/{document_id}/download"
    title = f"Document: {doc['title']}"
    body = (payload.message + "\n\n" if payload.message else "") + f"{doc['title']} ({doc['filename']})"
    email_html = _document_email_html(doc["title"], payload.message, abs_link, from_name)
    result = await deliver_broadcast(users, title, body, payload.channels, from_name, "document",
                                     email_html=email_html, link=rel_link, link_label="Download document",
                                     from_email=from_email)
    return result


@api_router.post("/learning-assignments")
async def create_learning_assignment(payload: LearningAssignmentCreate, staff: dict = Depends(require_staff)):
    doc = await db.documents.find_one({"id": payload.document_id}, {"_id": 0, "gridfs_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Workbook document not found")
    cadet_ids = list(dict.fromkeys([cid for cid in payload.cadet_ids if cid]))
    if not cadet_ids:
        raise HTTPException(status_code=400, detail="Select at least one cadet")
    cadet_count = await db.users.count_documents({"id": {"$in": cadet_ids}, "role": "cadet"})
    if cadet_count != len(cadet_ids):
        raise HTTPException(status_code=400, detail="Assignments can only be sent to cadets")

    assignment = {
        "id": str(uuid.uuid4()),
        "document_id": payload.document_id,
        "cadet_ids": cadet_ids,
        "instructions": payload.instructions.strip(),
        "due_date": (payload.due_date or "").strip(),
        "created_by": staff["id"],
        "created_by_name": f"{staff.get('first_name', '')} {staff.get('last_name', '')}".strip() or "Staff",
        "created_at": now_iso(),
    }
    await db.learning_assignments.insert_one(assignment)

    cadets = await db.users.find({"id": {"$in": cadet_ids}}, {"_id": 0, "id": 1, "first_name": 1, "last_name": 1}).to_list(500)
    notes = []
    for cadet in cadets:
        body = payload.instructions.strip() or "Open Documents to download your workbook and submit it when complete."
        notes.append({
            "id": str(uuid.uuid4()),
            "user_id": cadet["id"],
            "title": f"New learning assigned: {doc['title']}",
            "body": body,
            "from_name": assignment["created_by_name"],
            "kind": "learning_assignment",
            "link": "/portal",
            "link_label": "Open portal",
            "channels": ["dashboard"],
            "read": False,
            "created_at": now_iso(),
        })
    if notes:
        await db.notifications.insert_many(notes)
        for note in notes:
            await push_to_user(note["user_id"], note["title"], note["body"], "/portal")

    return await _learning_assignment_out(assignment)


@api_router.get("/learning-assignments/staff")
async def list_learning_assignments_staff(staff: dict = Depends(require_staff)):
    rows = await db.learning_assignments.find({}, {"_id": 0}).sort("created_at", -1).to_list(500)
    return [await _learning_assignment_out(r) for r in rows]


@api_router.get("/learning-assignments/my")
async def list_my_learning_assignments(user: dict = Depends(get_current_user)):
    if user["role"] == "cadet":
        rows = await db.learning_assignments.find({"cadet_ids": user["id"]}, {"_id": 0}).sort("created_at", -1).to_list(500)
        return [await _learning_assignment_out(r, viewer=user) for r in rows]
    if user["role"] == "parent":
        linked = user.get("child_ids", []) or []
        if not linked:
            return []
        rows = await db.learning_assignments.find({"cadet_ids": {"$in": linked}}, {"_id": 0}).sort("created_at", -1).to_list(500)
        return [await _learning_assignment_out(r) for r in rows]
    return []


@api_router.post("/learning-assignments/{assignment_id}/submit")
async def submit_learning_assignment(
        assignment_id: str, file: UploadFile = File(...), user: dict = Depends(get_current_user)):
    if user["role"] != "cadet":
        raise HTTPException(status_code=403, detail="Cadets only")
    assignment = await db.learning_assignments.find_one({"id": assignment_id}, {"_id": 0})
    if not assignment:
        raise HTTPException(status_code=404, detail="Learning assignment not found")
    if user["id"] not in (assignment.get("cadet_ids", []) or []):
        raise HTTPException(status_code=403, detail="This assignment is not for you")
    data = await file.read()
    if len(data) > 15 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="File too large (maximum 15 MB).")

    existing = await db.learning_submissions.find_one({"assignment_id": assignment_id, "cadet_id": user["id"]}, {"_id": 0})
    fid = existing.get("id") if existing else str(uuid.uuid4())
    gid = await fs.upload_from_stream(
        file.filename or "submission",
        data,
        metadata={"learning_submission_id": fid, "content_type": file.content_type},
    )
    if existing:
        try:
            await fs.delete(ObjectId(existing["gridfs_id"]))
        except Exception:
            pass

    submission = {
        "id": fid,
        "assignment_id": assignment_id,
        "cadet_id": user["id"],
        "filename": file.filename or "submission",
        "content_type": file.content_type or "application/octet-stream",
        "size": len(data),
        "gridfs_id": str(gid),
        "submitted_at": now_iso(),
    }
    await db.learning_submissions.update_one(
        {"assignment_id": assignment_id, "cadet_id": user["id"]},
        {"$set": submission},
        upsert=True,
    )

    workbook = await db.documents.find_one({"id": assignment.get("document_id")}, {"_id": 0, "title": 1})
    cadet_name = f"{user.get('first_name', '')} {user.get('last_name', '')}".strip() or "Cadet"
    reviewers = await _learning_reviewer_ids()
    notifications = []
    for reviewer_id in reviewers:
        notifications.append({
            "id": str(uuid.uuid4()),
            "user_id": reviewer_id,
            "title": f"Workbook submitted: {workbook.get('title', 'Learning assignment') if workbook else 'Learning assignment'}",
            "body": f"{cadet_name} has uploaded completed work for marking.",
            "from_name": cadet_name,
            "kind": "learning_submission",
            "link": f"/api/learning-submissions/{submission['id']}/download",
            "link_label": "Download submission",
            "channels": ["dashboard"],
            "read": False,
            "created_at": now_iso(),
        })
    if notifications:
        await db.notifications.insert_many(notifications)
        for note in notifications:
            await push_to_user(note["user_id"], note["title"], note["body"], "/portal")

    return {k: v for k, v in submission.items() if k != "gridfs_id"}


@api_router.get("/learning-submissions/{submission_id}/download")
async def download_learning_submission(submission_id: str, user: dict = Depends(get_current_user)):
    submission = await db.learning_submissions.find_one({"id": submission_id}, {"_id": 0})
    if not submission:
        raise HTTPException(status_code=404, detail="Submission not found")
    if user["role"] == "cadet":
        if submission.get("cadet_id") != user["id"]:
            raise HTTPException(status_code=403, detail="Forbidden")
    elif user["role"] == "parent":
        linked = user.get("child_ids", []) or []
        if submission.get("cadet_id") not in linked:
            raise HTTPException(status_code=403, detail="Forbidden")
    elif user["role"] not in ("admin", "cfav"):
        raise HTTPException(status_code=403, detail="Forbidden")
    stream = await fs.open_download_stream(ObjectId(submission["gridfs_id"]))
    data = await stream.read()
    return Response(content=data, media_type=submission["content_type"],
                    headers={"Content-Disposition": f'inline; filename="{submission["filename"]}"'})


# ---------------------------------------------------------------------------
# Blog / News (staff author; published posts are public + shown in portal)
# ---------------------------------------------------------------------------
def _slugify(text: str) -> str:
    s = re.sub(r"[^a-z0-9]+", "-", (text or "").lower()).strip("-")
    return (s or "post")[:60] + "-" + uuid.uuid4().hex[:6]


class BlogCreate(BaseModel):
    title: str = Field(..., min_length=2, max_length=200)
    excerpt: str = Field(default="", max_length=500)
    body: str = Field(..., min_length=1, max_length=40000)
    cover_image_url: str = Field(default="", max_length=1000)
    images: List[str] = []
    facebook_post_url: str = Field(default="", max_length=2000)
    status: str = "draft"  # draft | published
    model_config = ConfigDict(extra="ignore")


def _blog_out(b: dict) -> dict:
    return {k: v for k, v in b.items() if k != "_id"}


def _blog_public(b: dict) -> dict:
    return {"title": b["title"], "slug": b["slug"], "excerpt": b.get("excerpt", ""),
            "cover_image_url": b.get("cover_image_url", ""), "images": b.get("images", []),
            "body": b.get("body", ""), "author_name": b.get("author_name", ""),
            "published_at": b.get("published_at"),
            "facebook_post_url": b.get("facebook_post_url", "")}


@api_router.post("/blogs")
async def create_blog(payload: BlogCreate, staff: dict = Depends(require_staff)):
    now = now_iso()
    b = {"id": str(uuid.uuid4()), "slug": _slugify(payload.title),
         "title": payload.title, "excerpt": payload.excerpt, "body": payload.body,
            "cover_image_url": payload.cover_image_url, "images": payload.images,
            "facebook_post_url": payload.facebook_post_url,
         "status": payload.status if payload.status in ("draft", "published") else "draft",
         "author_id": staff["id"],
         "author_name": f"{staff.get('first_name','')} {staff.get('last_name','')}".strip() or "Staff",
         "created_at": now, "updated_at": now,
         "published_at": now if payload.status == "published" else None}
    await db.blogs.insert_one(b)
    return _blog_out(b)


@api_router.get("/blogs")
async def list_blogs(staff: dict = Depends(require_staff)):
    rows = await db.blogs.find({}, {"_id": 0}).sort("created_at", -1).to_list(500)
    return rows


@api_router.patch("/blogs/{blog_id}")
async def update_blog(blog_id: str, payload: BlogCreate, staff: dict = Depends(require_staff)):
    existing = await db.blogs.find_one({"id": blog_id})
    if not existing:
        raise HTTPException(status_code=404, detail="Post not found")
    upd = {"title": payload.title, "excerpt": payload.excerpt, "body": payload.body,
            "cover_image_url": payload.cover_image_url, "images": payload.images,
            "facebook_post_url": payload.facebook_post_url,
           "status": payload.status if payload.status in ("draft", "published") else "draft",
           "updated_at": now_iso()}
    if payload.status == "published" and not existing.get("published_at"):
        upd["published_at"] = now_iso()
    if payload.status == "draft":
        upd["published_at"] = None
    b = await db.blogs.find_one_and_update({"id": blog_id}, {"$set": upd},
                                           projection={"_id": 0}, return_document=True)
    return b


@api_router.delete("/blogs/{blog_id}")
async def delete_blog(blog_id: str, staff: dict = Depends(require_staff)):
    res = await db.blogs.delete_one({"id": blog_id})
    if res.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Post not found")
    return {"deleted": True}


@api_router.post("/blogs/upload-image")
async def upload_blog_image(file: UploadFile = File(...), staff: dict = Depends(require_staff)):
    data = await file.read()
    if len(data) > 8 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="Image too large (maximum 8MB).")
    iid = str(uuid.uuid4())
    gid = await fs.upload_from_stream(file.filename or "image", data,
                                      metadata={"content_type": file.content_type, "blog_image_id": iid})
    await db.blog_images.insert_one({"id": iid, "gridfs_id": str(gid),
                                     "content_type": file.content_type or "image/jpeg", "created_at": now_iso()})
    return {"id": iid, "url": f"/api/blogs/image/{iid}"}


@api_router.get("/blogs/image/{image_id}")
async def get_blog_image(image_id: str):
    img = await db.blog_images.find_one({"id": image_id}, {"_id": 0})
    if not img:
        raise HTTPException(status_code=404, detail="Image not found")
    stream = await fs.open_download_stream(ObjectId(img["gridfs_id"]))
    data = await stream.read()
    return Response(content=data, media_type=img["content_type"],
                    headers={"Cache-Control": "public, max-age=86400"})


@api_router.get("/public/blogs")
async def public_blogs():
    rows = await db.blogs.find({"status": "published"}, {"_id": 0}).sort("published_at", -1).to_list(200)
    return [{"title": b["title"], "slug": b["slug"], "excerpt": b.get("excerpt", ""),
             "cover_image_url": b.get("cover_image_url", ""), "author_name": b.get("author_name", ""),
             "published_at": b.get("published_at"), "facebook_post_url": b.get("facebook_post_url", "")}
            for b in rows]


@api_router.get("/public/blogs/{slug}")
async def public_blog(slug: str):
    b = await db.blogs.find_one({"slug": slug, "status": "published"}, {"_id": 0})
    if not b:
        raise HTTPException(status_code=404, detail="Post not found")
    return _blog_public(b)


# ---------------------------------------------------------------------------
# User management (staff)
# ---------------------------------------------------------------------------
@api_router.get("/users")
async def list_users(staff: dict = Depends(require_privileged_staff), role: Optional[str] = None,
                     uniformed: Optional[bool] = None):
    q = {"role": role} if role else {}
    if role == "cfav" and uniformed is not None:
        q["is_uniformed"] = bool(uniformed)
    users = await db.users.find(q, {"_id": 0}).sort("first_name", 1).to_list(2000)
    out = []
    for u in users:
        pu = public_user(u)
        if u["role"] in ("cadet",):
            pu["stats"] = await compute_member_stats(u["id"])
        out.append(pu)
    return out


@api_router.post("/users")
async def create_user(payload: UserCreate, staff: dict = Depends(require_privileged_staff)):
    if payload.role not in ALL_ROLES:
        raise HTTPException(status_code=400, detail="Invalid role")
    email = (str(payload.email).strip().lower() if payload.email else "")
    if payload.role != "cadet" and not email:
        raise HTTPException(status_code=400, detail="Email is required for this role")
    if payload.role != "cadet" and len((payload.password or "")) < 6:
        raise HTTPException(status_code=400, detail="Password must be at least 6 characters")
    if email:
        existing = await db.users.find_one({"email": email})
        if existing:
            raise HTTPException(status_code=400, detail="A user with this email already exists")
    username = await ensure_login_username(payload.first_name, payload.last_name)
    chosen_password = CADET_DEFAULT_PASSWORD if payload.role == "cadet" else payload.password
    user = {
        "id": str(uuid.uuid4()), "email": email,
        "password_hash": hash_password(chosen_password), "role": payload.role,
        "login_username": username,
        "must_change_password": payload.role == "cadet",
        "first_name": payload.first_name, "last_name": payload.last_name,
        "is_uniformed": bool(payload.is_uniformed) if payload.role == "cfav" else None,
        "child_ids": payload.child_ids if payload.role == "parent" else [],
        "bonus_points": 0, "created_at": now_iso(),
    }
    await db.users.insert_one(user)
    return public_user(user)


@api_router.post("/users/import-upload")
async def import_users_upload(
        file: UploadFile = File(...),
        role_hint: str = Form(""),
        staff: dict = Depends(require_privileged_staff)):
    data = await file.read()
    if len(data) > 10 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="File too large (maximum 10MB).")

    filename = (file.filename or "").lower()
    records: List[dict]
    try:
        if filename.endswith(".csv"):
            import pandas as pd
            df = pd.read_csv(io.BytesIO(data))
            records = _extract_registration_rows(df.fillna("").to_dict(orient="records"), role_hint=role_hint)
        elif filename.endswith(".xlsx") or filename.endswith(".xls"):
            import pandas as pd
            df = pd.read_excel(io.BytesIO(data))
            records = _extract_registration_rows(df.fillna("").to_dict(orient="records"), role_hint=role_hint)
        elif filename.endswith(".docx"):
            records = _parse_docx_registration_rows(data, role_hint=role_hint)
        else:
            raise HTTPException(status_code=400, detail="Please upload .xlsx, .xls, .csv or .docx.")
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Could not read file: {exc}")

    if not records:
        raise HTTPException(status_code=400, detail="No valid cadet/CFAV rows found.")

    created, skipped, errors = [], [], []
    for idx, rec in enumerate(records, start=1):
        role = rec.get("role", "")
        first_name = (rec.get("first_name") or "").strip()
        last_name = (rec.get("last_name") or "").strip()
        email = (rec.get("email") or "").strip().lower()
        is_uniformed = bool(rec.get("is_uniformed", False))
        try:
            if role not in {"cadet", "cfav"}:
                errors.append({"row": idx, "reason": "Role must be cadet or cfav"})
                continue
            if role == "cfav":
                if not email:
                    errors.append({"row": idx, "reason": "CFAV row requires RAFAC email"})
                    continue
                if not _is_rafac_email(email):
                    errors.append({"row": idx, "reason": "CFAV email must end @rafac.mod.gov.uk"})
                    continue

            existing = await _find_registration_duplicate(role, first_name, last_name, email)
            if existing:
                skipped.append({
                    "row": idx,
                    "reason": "Already exists",
                    "name": f"{existing.get('first_name', '')} {existing.get('last_name', '')}".strip(),
                    "role": existing.get("role", role),
                    "login_username": existing.get("login_username", ""),
                })
                continue

            user = await _create_registered_user(role, first_name, last_name, email, is_uniformed)
            created.append({
                "id": user["id"],
                "role": user["role"],
                "name": f"{user.get('first_name', '')} {user.get('last_name', '')}".strip(),
                "login_username": user.get("login_username", ""),
                "email": user.get("email", ""),
                "password": CADET_DEFAULT_PASSWORD,
                "must_change_password": True,
            })
        except Exception as exc:
            errors.append({"row": idx, "reason": str(exc)})

    return {
        "created": len(created),
        "skipped": len(skipped),
        "errors": len(errors),
        "created_users": created,
        "skipped_rows": skipped,
        "error_rows": errors,
        "default_password": CADET_DEFAULT_PASSWORD,
    }


@api_router.get("/users/import-template")
async def users_import_template(
        format: str = "xlsx",
        role: str = "cadet",
        staff: dict = Depends(require_privileged_staff)):
    role_key = _normalise_role(role)
    if role_key not in {"cadet", "cfav"}:
        raise HTTPException(status_code=400, detail="role must be cadet or cfav")

    rows = [
        {
            "role": role_key,
            "first_name": "Sam",
            "last_name": "Cadet" if role_key == "cadet" else "Instructor",
            "email": "" if role_key == "cadet" else "sam.instructor@rafac.mod.gov.uk",
            "is_uniformed": "" if role_key == "cadet" else "yes",
        },
        {
            "role": role_key,
            "first_name": "Alex",
            "last_name": "Example",
            "email": "" if role_key == "cadet" else "alex.example@rafac.mod.gov.uk",
            "is_uniformed": "" if role_key == "cadet" else "no",
        },
    ]

    fmt = (format or "xlsx").strip().lower()
    if fmt in {"xlsx", "excel"}:
        import pandas as pd

        df = pd.DataFrame(rows, columns=["role", "first_name", "last_name", "email", "is_uniformed"])
        out = io.BytesIO()
        with pd.ExcelWriter(out, engine="openpyxl") as writer:
            df.to_excel(writer, sheet_name="members", index=False)
        body = out.getvalue()
        name = f"member-import-template-{role_key}.xlsx"
        return Response(
            content=body,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f'attachment; filename="{name}"'},
        )

    if fmt in {"docx", "word"}:
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_heading("1471 Member Import Template", level=1)
        doc.add_paragraph("Complete the table, then upload this file in Members > Import Excel/CSV/Word.")
        table = doc.add_table(rows=1, cols=5)
        heads = ["role", "first_name", "last_name", "email", "is_uniformed"]
        for idx, head in enumerate(heads):
            table.rows[0].cells[idx].text = head
        for row in rows:
            cells = table.add_row().cells
            cells[0].text = row["role"]
            cells[1].text = row["first_name"]
            cells[2].text = row["last_name"]
            cells[3].text = row["email"]
            cells[4].text = row["is_uniformed"]
        out = io.BytesIO()
        doc.save(out)
        body = out.getvalue()
        name = f"member-import-template-{role_key}.docx"
        return Response(
            content=body,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{name}"'},
        )

    raise HTTPException(status_code=400, detail="format must be xlsx or docx")


@api_router.get("/users/register-qr")
async def users_register_qr(
        role: str = "cadet",
        base_url: str = "",
        staff: dict = Depends(require_privileged_staff)):
    role_key = _normalise_role(role)
    if role_key not in {"cadet", "cfav"}:
        raise HTTPException(status_code=400, detail="role must be cadet or cfav")

    root = (base_url or "").strip()
    if not root:
        root = str(os.environ.get("PUBLIC_APP_URL") or os.environ.get("CORS_ORIGINS", "").split(",")[0]).strip()
    if not root:
        root = "https://1471squadron.co.uk"
    root = root.rstrip("/")
    if not root.startswith("http://") and not root.startswith("https://"):
        root = f"https://{root}"
    target = f"{root}/register?role={role_key}"

    try:
        import qrcode
    except Exception:
        raise HTTPException(status_code=500, detail="QR generator is unavailable on this server")

    img = qrcode.make(target)
    out = io.BytesIO()
    img.save(out, format="PNG")
    filename = f"register-{role_key}-qr.png"
    return Response(
        content=out.getvalue(),
        media_type="image/png",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Cache-Control": "no-store",
        },
    )


@api_router.post("/public/register-self")
async def register_self(payload: SelfRegisterRequest):
    role = _normalise_role(payload.role)
    if role not in {"cadet", "cfav"}:
        raise HTTPException(status_code=400, detail="Role must be cadet or cfav")

    first_name = payload.first_name.strip()
    last_name = payload.last_name.strip()
    email = (str(payload.email).strip().lower() if payload.email else "")
    if role == "cfav":
        if not email:
            raise HTTPException(status_code=400, detail="RAFAC email is required for CFAV registration")
        if not _is_rafac_email(email):
            raise HTTPException(status_code=400, detail="CFAV email must end @rafac.mod.gov.uk")

    existing = await _find_registration_duplicate(role, first_name, last_name, email)
    if existing:
        raise HTTPException(status_code=400, detail="An account for these details already exists")

    user = await _create_registered_user(role, first_name, last_name, email, bool(payload.is_uniformed))
    return {
        "created": True,
        "role": user["role"],
        "login_username": user.get("login_username", ""),
        "email": user.get("email", ""),
        "default_password": CADET_DEFAULT_PASSWORD,
        "must_change_password": True,
    }


@api_router.patch("/users/{user_id}")
async def update_user(user_id: str, payload: UserUpdate, staff: dict = Depends(require_privileged_staff)):
    updates = {k: v for k, v in payload.model_dump().items() if v is not None}
    if "role" in updates and updates["role"] != "cfav":
        updates["is_uniformed"] = None
    if "is_uniformed" in updates and "role" not in updates:
        existing = await db.users.find_one({"id": user_id}, {"_id": 0, "role": 1})
        if not existing:
            raise HTTPException(status_code=404, detail="User not found")
        if existing.get("role") != "cfav":
            updates["is_uniformed"] = None
    if "role" in updates and updates["role"] not in ALL_ROLES:
        raise HTTPException(status_code=400, detail="Invalid role")
    if "first_name" in updates or "last_name" in updates:
        existing = await db.users.find_one({"id": user_id}, {"_id": 0, "first_name": 1, "last_name": 1})
        if not existing:
            raise HTTPException(status_code=404, detail="User not found")
        new_first = updates.get("first_name", existing.get("first_name", ""))
        new_last = updates.get("last_name", existing.get("last_name", ""))
        updates["login_username"] = await ensure_login_username(new_first, new_last, user_id)
    if not updates:
        raise HTTPException(status_code=400, detail="No fields to update")
    res = await db.users.find_one_and_update(
        {"id": user_id}, {"$set": updates}, projection={"_id": 0, "password_hash": 0},
        return_document=True)
    if not res:
        raise HTTPException(status_code=404, detail="User not found")
    return public_user(res)


@api_router.post("/users/{user_id}/reset-password")
async def reset_user_password(user_id: str, payload: ResetPassword, staff: dict = Depends(require_privileged_staff)):
    res = await db.users.update_one({"id": user_id},
                                    {"$set": {"password_hash": hash_password(payload.new_password),
                                              "must_change_password": True}})
    if res.matched_count == 0:
        raise HTTPException(status_code=404, detail="User not found")
    return {"updated": True}


@api_router.post("/users/{user_id}/reset-cadet-password")
async def reset_cadet_password(user_id: str, staff: dict = Depends(require_privileged_staff)):
    target = await db.users.find_one({"id": user_id}, {"_id": 0, "role": 1})
    if not target:
        raise HTTPException(status_code=404, detail="User not found")
    if target.get("role") != "cadet":
        raise HTTPException(status_code=400, detail="This reset is for cadet accounts only")
    await db.users.update_one(
        {"id": user_id},
        {"$set": {"password_hash": hash_password(CADET_DEFAULT_PASSWORD), "must_change_password": True}},
    )
    return {"updated": True, "default_password": CADET_DEFAULT_PASSWORD}


@api_router.delete("/users/{user_id}")
async def delete_user(user_id: str, staff: dict = Depends(require_privileged_staff)):
    target = await db.users.find_one({"id": user_id})
    if not target:
        raise HTTPException(status_code=404, detail="User not found")
    await db.users.delete_one({"id": user_id})
    return {"deleted": True}


@api_router.get("/appointments")
async def get_appointments(staff: dict = Depends(require_roles("admin"))):
    doc = await _appointments_doc()
    value = doc.get("value", {})
    out = {}
    for k in APPOINTMENT_KEYS:
        uid = value.get(k)
        if not uid:
            out[k] = None
            continue
        u = await db.users.find_one({"id": uid}, {"_id": 0, "password_hash": 0})
        out[k] = public_user(u) if u else None
    return out


@api_router.put("/appointments")
async def set_appointments(payload: AppointmentsUpdate, staff: dict = Depends(require_roles("admin"))):
    val = payload.model_dump()
    cleaned = {}
    for k, uid in val.items():
        if k not in APPOINTMENT_KEYS:
            continue
        if uid in (None, ""):
            cleaned[k] = None
            continue
        u = await db.users.find_one({"id": uid}, {"_id": 0})
        if not u or u.get("role") not in ("admin", "cfav"):
            raise HTTPException(status_code=400, detail=f"Invalid appointment user for {k}")
        cleaned[k] = uid

    await db.settings.update_one(
        {"key": "appointments"},
        {"$set": {"key": "appointments", "value": cleaned, "updated_at": now_iso(), "updated_by": staff["id"]}},
        upsert=True,
    )
    return {"saved": True}


# ---------------------------------------------------------------------------
# Cadet status tracker (Excel import + progression alerts)
# ---------------------------------------------------------------------------
def _pick_col(cols: List[str], candidates: List[str]) -> Optional[str]:
    folded = {re.sub(r"[^a-z0-9]+", "", c.lower()): c for c in cols}
    for cand in candidates:
        key = re.sub(r"[^a-z0-9]+", "", cand.lower())
        if key in folded:
            return folded[key]
    return None


def _badge_list(val) -> List[str]:
    if val is None:
        return []
    s = str(val).strip()
    if not s:
        return []
    parts = re.split(r"[,;|/]", s)
    out = [p.strip() for p in parts if p.strip()]
    seen, uniq = set(), []
    for x in out:
        k = x.lower()
        if k in seen:
            continue
        seen.add(k)
        uniq.append(x)
    return uniq[:20]


def _normalise_role(value: str) -> str:
    v = (value or "").strip().lower()
    if v in {"cadet", "c"}:
        return "cadet"
    if v in {"cfav", "staff", "adult", "volunteer"}:
        return "cfav"
    if v in {"admin", "parent"}:
        return v
    return ""


def _to_bool(value) -> Optional[bool]:
    if isinstance(value, bool):
        return value
    s = str(value or "").strip().lower()
    if not s:
        return None
    if s in {"1", "true", "yes", "y", "uniformed", "u", "commissioned", "ci", "snp"}:
        return True
    if s in {"0", "false", "no", "n", "non-uniformed", "non uniformed", "civilian", "none"}:
        return False
    return None


def _is_rafac_email(email: str) -> bool:
    return (email or "").strip().lower().endswith("@rafac.mod.gov.uk")


def _extract_registration_rows(rows: List[dict], role_hint: str = "") -> List[dict]:
    if not rows:
        return []
    cols = [str(c) for c in rows[0].keys()]
    role_col = _pick_col(cols, ["role", "member type", "type", "account type"])
    first_col = _pick_col(cols, ["first name", "firstname", "given name", "forename"])
    last_col = _pick_col(cols, ["last name", "lastname", "surname", "family name"])
    full_name_col = _pick_col(cols, ["name", "full name", "cadet name", "member name"])
    email_col = _pick_col(cols, ["email", "email address", "rafac email", "staff email"])
    uniformed_col = _pick_col(cols, ["is uniformed", "uniformed", "uniform", "cfav uniformed"])

    out = []
    hint = _normalise_role(role_hint)
    for row in rows:
        role = _normalise_role(str(row.get(role_col, ""))) if role_col else ""
        role = role or hint
        if role not in {"cadet", "cfav"}:
            continue

        first = str(row.get(first_col, "")).strip() if first_col else ""
        last = str(row.get(last_col, "")).strip() if last_col else ""
        if (not first or not last) and full_name_col:
            full = str(row.get(full_name_col, "")).strip()
            if full:
                parts = full.split()
                if parts and not first:
                    first = parts[0]
                if len(parts) > 1 and not last:
                    last = " ".join(parts[1:])
        if not first:
            continue

        email = str(row.get(email_col, "")).strip().lower() if email_col else ""
        is_uniformed = _to_bool(row.get(uniformed_col, "")) if uniformed_col else None
        out.append({
            "role": role,
            "first_name": first,
            "last_name": last,
            "email": email,
            "is_uniformed": bool(is_uniformed) if is_uniformed is not None else False,
        })
    return out


def _parse_docx_registration_rows(data: bytes, role_hint: str = "") -> List[dict]:
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(data))
    rows = []
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            table_rows.append([c.text.strip() for c in row.cells])
        if len(table_rows) < 2:
            continue
        headers = [h.strip() for h in table_rows[0]]
        for vals in table_rows[1:]:
            rec = {}
            for idx, head in enumerate(headers):
                if head:
                    rec[head] = vals[idx].strip() if idx < len(vals) else ""
            if rec:
                rows.append(rec)
    return _extract_registration_rows(rows, role_hint=role_hint)


async def _find_registration_duplicate(role: str, first_name: str, last_name: str, email: str) -> Optional[dict]:
    if email:
        existing = await db.users.find_one({"email": email}, {"_id": 0})
        if existing:
            return existing
    if role == "cadet":
        return await db.users.find_one({
            "role": "cadet",
            "first_name": {"$regex": f"^{re.escape(first_name)}$", "$options": "i"},
            "last_name": {"$regex": f"^{re.escape(last_name)}$", "$options": "i"},
        }, {"_id": 0})
    return None


async def _create_registered_user(role: str, first_name: str, last_name: str, email: str, is_uniformed: bool) -> dict:
    username = await ensure_login_username(first_name, last_name)
    stored_email = (email or "").strip().lower()
    if role == "cadet" and not stored_email:
        stored_email = f"{username}@cadet.local"

    user = {
        "id": str(uuid.uuid4()),
        "email": stored_email,
        "password_hash": hash_password(CADET_DEFAULT_PASSWORD),
        "role": role,
        "login_username": username,
        "must_change_password": True,
        "first_name": first_name,
        "last_name": last_name,
        "is_uniformed": bool(is_uniformed) if role == "cfav" else None,
        "child_ids": [],
        "bonus_points": 0,
        "created_at": now_iso(),
    }
    await db.users.insert_one(user)
    return user


async def _send_classification_stagnation_alerts() -> dict:
    rows = await db.cadet_tracker.find({}, {"_id": 0}).to_list(5000)
    if not rows:
        return {"processed": 0, "cadet_messages": 0, "admin_notifications": 0}

    admins = await db.users.find({"role": "admin"}, {"_id": 0, "id": 1}).to_list(200)
    now = datetime.now(timezone.utc)
    cadet_messages = 0
    admin_notifications = 0

    for r in rows:
        cadet_id = r.get("user_id")
        if not cadet_id:
            continue
        changed_at = r.get("classification_changed_at") or r.get("updated_at") or r.get("created_at")
        if not changed_at:
            continue
        try:
            last_dt = datetime.fromisoformat(changed_at.replace("Z", "+00:00"))
        except Exception:
            continue
        days = (now - last_dt).days
        if days < 30:
            continue
        level = 3 if days >= 90 else (2 if days >= 60 else 1)

        state = await db.cadet_progress_alerts.find_one({"cadet_id": cadet_id}, {"_id": 0})
        sent = set((state or {}).get("sent_levels", []))
        if level in sent:
            continue

        # Ensure we send missing lower levels once as cadence catches up.
        to_send = [lv for lv in (1, 2, 3) if lv <= level and lv not in sent]
        for lv in to_send:
            body = (
                "Your classification has not changed for 1 month. Speak with staff about your next classification target."
                if lv == 1 else
                "Your classification has not changed for 2 months. Please ask staff for a progression plan."
                if lv == 2 else
                "Your classification has not changed for 3 months. Staff have been notified to help plan support."
            )
            msg = {
                "id": str(uuid.uuid4()), "member_id": cadet_id,
                "member_name": r.get("name", "Cadet"),
                "from_staff": True, "author": "Training Team",
                "body": body, "created_at": now_iso(),
                "read_by_member": False, "read_by_staff": True,
            }
            await db.messages.insert_one(msg)
            await push_to_user(cadet_id, "Classification progression", body, "/portal")
            cadet_messages += 1

            if lv == 3 and admins:
                notes = [{
                    "id": str(uuid.uuid4()),
                    "user_id": a["id"],
                    "title": "3-month classification stall",
                    "body": f"{r.get('name','Cadet')} has had no classification change for 3 months.",
                    "from_name": "Training Monitor",
                    "kind": "progression_alert",
                    "channels": ["dashboard"],
                    "read": False,
                    "created_at": now_iso(),
                } for a in admins]
                await db.notifications.insert_many(notes)
                for n in notes:
                    await push_to_user(n["user_id"], n["title"], n["body"], "/portal")
                admin_notifications += len(notes)

        sent.update(to_send)
        await db.cadet_progress_alerts.update_one(
            {"cadet_id": cadet_id},
            {"$set": {
                "cadet_id": cadet_id,
                "sent_levels": sorted(sent),
                "classification": r.get("classification", ""),
                "classification_changed_at": r.get("classification_changed_at"),
                "updated_at": now_iso(),
            }},
            upsert=True,
        )

    return {
        "processed": len(rows),
        "cadet_messages": cadet_messages,
        "admin_notifications": admin_notifications,
    }


@api_router.post("/cadet-tracker/upload")
async def upload_cadet_tracker(file: UploadFile = File(...), staff: dict = Depends(require_roles("admin"))):
    data = await file.read()
    if len(data) > 10 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="File too large (maximum 10MB).")

    import pandas as pd
    try:
        if (file.filename or "").lower().endswith(".csv"):
            df = pd.read_csv(io.BytesIO(data))
        else:
            df = pd.read_excel(io.BytesIO(data))
    except Exception:
        raise HTTPException(status_code=400, detail="Could not read file. Upload .xlsx or .csv cadet tracker.")

    cols = [str(c) for c in df.columns]
    name_col = _pick_col(cols, ["name", "cadet name", "full name"])
    rank_col = _pick_col(cols, ["rank", "cadet rank"])
    class_col = _pick_col(cols, ["classification", "class", "training classification"])
    badge_col = _pick_col(cols, ["badges", "major badges", "achievement badges"])
    email_col = _pick_col(cols, ["email", "cadet email", "contact email"])
    if not name_col:
        raise HTTPException(status_code=400, detail="Tracker must include a cadet name column.")

    created, updated = 0, 0
    out_rows = []
    for _, row in df.fillna("").iterrows():
        nm = str(row.get(name_col, "")).strip()
        if not nm:
            continue
        rank = str(row.get(rank_col, "")).strip() if rank_col else ""
        classification = str(row.get(class_col, "")).strip() if class_col else ""
        badges = _badge_list(row.get(badge_col, "")) if badge_col else []
        em = str(row.get(email_col, "")).strip().lower() if email_col else ""

        cadet = None
        if em:
            cadet = await db.users.find_one({"email": em, "role": "cadet"}, {"_id": 0})
        if cadet is None:
            parts = nm.split()
            first = parts[0].lower() if parts else ""
            last = " ".join(parts[1:]).lower() if len(parts) > 1 else ""
            if first:
                cadet = await db.users.find_one({
                    "role": "cadet",
                    "first_name": {"$regex": f"^{re.escape(first)}$", "$options": "i"},
                    "last_name": {"$regex": f"^{re.escape(last)}$", "$options": "i"},
                }, {"_id": 0})

        tracker_key = cadet["id"] if cadet else nm.lower()
        existing = await db.cadet_tracker.find_one({"tracker_key": tracker_key}, {"_id": 0})
        changed_at = existing.get("classification_changed_at") if existing else None
        prev_class = (existing or {}).get("classification", "")
        if classification and classification != prev_class:
            changed_at = now_iso()
            await db.cadet_progress_alerts.delete_one({"cadet_id": (cadet or {}).get("id")})
        if not changed_at:
            changed_at = now_iso()

        doc = {
            "tracker_key": tracker_key,
            "user_id": cadet["id"] if cadet else None,
            "name": nm,
            "email": cadet.get("email") if cadet else em,
            "rank": rank,
            "classification": classification,
            "major_badges": badges,
            "classification_changed_at": changed_at,
            "updated_at": now_iso(),
            "updated_by": staff["id"],
        }
        res = await db.cadet_tracker.update_one({"tracker_key": tracker_key}, {"$set": doc}, upsert=True)
        if res.upserted_id:
            created += 1
        else:
            updated += 1
        out_rows.append(doc)

    return {"created": created, "updated": updated, "rows": out_rows}


@api_router.get("/cadet-tracker")
async def list_cadet_tracker(staff: dict = Depends(require_staff)):
    return await db.cadet_tracker.find({}, {"_id": 0}).sort("name", 1).to_list(5000)


@api_router.post("/cadet-tracker/run-alerts")
async def run_cadet_tracker_alerts(staff: dict = Depends(require_roles("admin"))):
    return await _send_classification_stagnation_alerts()


# ---------------------------------------------------------------------------
# CFAV planning inputs (availability, event ideas, skill matrix)
# ---------------------------------------------------------------------------
@api_router.post("/cfav/availability")
async def submit_cfav_availability(payload: CfavAvailabilityCreate, cfav: dict = Depends(require_roles("cfav"))):
    doc = {
        "id": str(uuid.uuid4()),
        "cfav_id": cfav["id"],
        "cfav_name": f"{cfav.get('first_name','')} {cfav.get('last_name','')}".strip() or "CFAV",
        "is_uniformed": bool(cfav.get("is_uniformed", False)),
        "parade_date": payload.parade_date,
        "available": bool(payload.available),
        "capabilities": payload.capabilities,
        "note": payload.note,
        "created_at": now_iso(),
    }
    await db.cfav_availability.insert_one(doc)
    return {k: v for k, v in doc.items() if k != "_id"}


@api_router.post("/cfav/event-ideas")
async def submit_cfav_event_idea(payload: CfavEventIdeaCreate, cfav: dict = Depends(require_roles("cfav"))):
    doc = {
        "id": str(uuid.uuid4()),
        "cfav_id": cfav["id"],
        "cfav_name": f"{cfav.get('first_name','')} {cfav.get('last_name','')}".strip() or "CFAV",
        "is_uniformed": bool(cfav.get("is_uniformed", False)),
        "title": payload.title,
        "parade_date": payload.parade_date,
        "summary": payload.summary,
        "created_at": now_iso(),
    }
    await db.cfav_event_ideas.insert_one(doc)
    return {k: v for k, v in doc.items() if k != "_id"}


@api_router.post("/cfav/skills")
async def submit_cfav_skills(payload: CfavSkillMatrixCreate, cfav: dict = Depends(require_roles("cfav"))):
    doc = {
        "cfav_id": cfav["id"],
        "cfav_name": f"{cfav.get('first_name','')} {cfav.get('last_name','')}".strip() or "CFAV",
        "is_uniformed": bool(cfav.get("is_uniformed", False)),
        "skills": payload.skills,
        "qualifications": payload.qualifications,
        "interested_activities": payload.interested_activities,
        "willing_to_support": payload.willing_to_support,
        "note": payload.note,
        "updated_at": now_iso(),
    }
    await db.cfav_skill_matrix.update_one({"cfav_id": cfav["id"]}, {"$set": doc}, upsert=True)
    return doc


@api_router.get("/cfav/skills/me")
async def get_cfav_skills_me(cfav: dict = Depends(require_roles("cfav"))):
    doc = await db.cfav_skill_matrix.find_one({"cfav_id": cfav["id"]}, {"_id": 0})
    return doc or {
        "cfav_id": cfav["id"],
        "skills": [],
        "qualifications": [],
        "interested_activities": [],
        "willing_to_support": [],
        "note": "",
    }


@api_router.get("/planning/cfav-inputs")
async def planning_cfav_inputs(staff: dict = Depends(require_staff), uniformed: Optional[bool] = None):
    fq = {"is_uniformed": bool(uniformed)} if uniformed is not None else {}
    av = await db.cfav_availability.find(fq, {"_id": 0}).sort("created_at", -1).to_list(5000)
    ideas = await db.cfav_event_ideas.find(fq, {"_id": 0}).sort("created_at", -1).to_list(5000)
    skills = await db.cfav_skill_matrix.find(fq, {"_id": 0}).sort("cfav_name", 1).to_list(5000)
    return {"availability": av, "event_ideas": ideas, "skills": skills}


# ---------------------------------------------------------------------------
# Training plan (month slots, CFAV bids, A4 export)
# ---------------------------------------------------------------------------
def _add_months(d: date, months: int) -> date:
    y = d.year + ((d.month - 1 + months) // 12)
    m = ((d.month - 1 + months) % 12) + 1
    return date(y, m, 1)


def _easter_sunday(year: int) -> date:
    # Anonymous Gregorian algorithm
    a = year % 19
    b = year // 100
    c = year % 100
    d = b // 4
    e = b % 4
    f = (b + 8) // 25
    g = (b - f + 1) // 3
    h = (19 * a + b - d - g + 15) % 30
    i = c // 4
    k = c % 4
    l = (32 + 2 * e + 2 * i - h - k) % 7
    m = (a + 11 * h + 22 * l) // 451
    month = (h + l - 7 * m + 114) // 31
    day = ((h + l - 7 * m + 114) % 31) + 1
    return date(year, month, day)


def _observed(dt: date) -> date:
    if dt.weekday() == 5:  # Saturday
        return dt + timedelta(days=2)
    if dt.weekday() == 6:  # Sunday
        return dt + timedelta(days=1)
    return dt


def _uk_bank_holidays_england_wales(year: int) -> set:
    easter = _easter_sunday(year)
    good_friday = easter - timedelta(days=2)
    easter_monday = easter + timedelta(days=1)

    new_year = _observed(date(year, 1, 1))
    early_may = date(year, 5, 1) + timedelta(days=(7 - date(year, 5, 1).weekday()) % 7)
    spring_bank = date(year, 5, 31) - timedelta(days=(date(year, 5, 31).weekday() - 0) % 7)
    summer_bank = date(year, 8, 31) - timedelta(days=(date(year, 8, 31).weekday() - 0) % 7)

    christmas = date(year, 12, 25)
    boxing = date(year, 12, 26)
    if christmas.weekday() == 5:  # Sat
        christmas_obs = date(year, 12, 27)
        boxing_obs = date(year, 12, 28)
    elif christmas.weekday() == 6:  # Sun
        christmas_obs = date(year, 12, 27)
        boxing_obs = date(year, 12, 26)
    elif boxing.weekday() == 5:  # Sat
        christmas_obs = christmas
        boxing_obs = date(year, 12, 28)
    elif boxing.weekday() == 6:  # Sun
        christmas_obs = christmas
        boxing_obs = date(year, 12, 27)
    else:
        christmas_obs = christmas
        boxing_obs = boxing

    return {
        new_year,
        good_friday,
        easter_monday,
        early_may,
        spring_bank,
        summer_bank,
        christmas_obs,
        boxing_obs,
    }


def _slot_out(s: dict, month_index: int, my_bid_count: int = 0) -> dict:
    return {
        "id": s["id"],
        "slot_date": s["slot_date"],
        "day_label": s.get("day_label", ""),
        "first_period_activity": s.get("first_period_activity", ""),
        "second_period_activity": s.get("second_period_activity", ""),
        "uniform_needed": s.get("uniform_needed", ""),
        "no_parade": bool(s.get("no_parade", False)),
        "no_parade_reason": s.get("no_parade_reason", ""),
        "month_index": month_index,
        "can_bid": month_index in (1, 2, 3) and not bool(s.get("no_parade", False)),
        "my_bid_count": my_bid_count,
    }


async def _ensure_training_slots(months_ahead: int = 3, actor_id: str = "system") -> int:
    today = date.today()
    created = 0
    for mo in range(1, max(1, months_ahead) + 1):
        target = _add_months(date(today.year, today.month, 1), mo)
        days = monthrange(target.year, target.month)[1]
        holidays = _uk_bank_holidays_england_wales(target.year)
        for dnum in range(1, days + 1):
            d = date(target.year, target.month, dnum)
            if d.weekday() not in (0, 3):
                continue
            is_holiday = d in holidays
            slot_id = f"tp-{d.isoformat()}"
            doc = {
                "id": slot_id,
                "slot_date": d.isoformat(),
                "day_label": d.strftime("%a"),
                "first_period_activity": "" if not is_holiday else "NO PARADE",
                "second_period_activity": "" if not is_holiday else "Bank holiday",
                "uniform_needed": "" if not is_holiday else "N/A",
                "no_parade": bool(is_holiday),
                "no_parade_reason": "Bank holiday" if is_holiday else "",
                "updated_at": now_iso(),
                "updated_by": actor_id,
            }
            res = await db.training_plan.update_one(
                {"slot_date": d.isoformat()},
                {"$setOnInsert": doc},
                upsert=True,
            )
            if res.upserted_id:
                created += 1
    return created


@api_router.post("/training-plan/populate-next-month")
async def populate_training_plan_next_month(staff: dict = Depends(require_privileged_staff)):
    created = await _ensure_training_slots(months_ahead=3, actor_id=staff["id"])
    target = _add_months(date.today().replace(day=1), 1)
    return {"month": target.strftime("%Y-%m"), "created": created, "months_seeded": 3}


@api_router.get("/training-plan")
async def list_training_plan(user: dict = Depends(require_roles("admin", "cfav")), months: int = 3):
    m = max(1, min(6, int(months or 3)))
    await _ensure_training_slots(months_ahead=max(3, m), actor_id="system")
    today = date.today()
    start = date(today.year, today.month, 1)
    end = _add_months(start, m + 1)
    rows = await db.training_plan.find({
        "slot_date": {"$gte": start.isoformat(), "$lt": end.isoformat()}
    }, {"_id": 0}).sort("slot_date", 1).to_list(2000)

    out = []
    for r in rows:
        sd = date.fromisoformat(r["slot_date"])
        month_idx = (sd.year - today.year) * 12 + (sd.month - today.month)
        my_count = 0
        if user["role"] == "cfav":
            my_count = await db.training_plan_bids.count_documents({"slot_id": r["id"], "cfav_id": user["id"]})
        out.append(_slot_out(r, month_idx, my_count))
    return out


@api_router.patch("/training-plan/{slot_id}")
async def update_training_slot(slot_id: str, payload: TrainingPlanSlotUpdate,
                               staff: dict = Depends(require_privileged_staff)):
    updates = payload.model_dump()
    updates["updated_at"] = now_iso()
    updates["updated_by"] = staff["id"]
    res = await db.training_plan.find_one_and_update(
        {"id": slot_id}, {"$set": updates}, projection={"_id": 0}, return_document=True)
    if not res:
        raise HTTPException(status_code=404, detail="Training slot not found")
    sd = date.fromisoformat(res["slot_date"])
    today = date.today()
    month_idx = (sd.year - today.year) * 12 + (sd.month - today.month)
    return _slot_out(res, month_idx)


@api_router.post("/training-plan/publish-month")
async def publish_training_plan_month(month: str, staff: dict = Depends(require_privileged_staff)):
    if not re.match(r"^\d{4}-\d{2}$", month or ""):
        raise HTTPException(status_code=400, detail="Month must be YYYY-MM")
    rows = await db.training_plan.find({
        "slot_date": {"$gte": f"{month}-01", "$lt": f"{month}-32"}
    }, {"_id": 0}).sort("slot_date", 1).to_list(500)

    created = 0
    updated = 0
    removed = 0
    for r in rows:
        slot_id = r.get("id")
        if not slot_id:
            continue
        slot_date = r["slot_date"]
        no_parade = bool(r.get("no_parade", False))
        p1 = (r.get("first_period_activity", "") or "").strip()
        p2 = (r.get("second_period_activity", "") or "").strip()
        uniform = (r.get("uniform_needed", "") or "").strip()

        # If no activity is planned or marked no parade, clear any previously published event for this slot.
        if no_parade or (not p1 and not p2):
            res = await db.events.delete_one({"training_slot_id": slot_id, "source": "training_plan"})
            if res.deleted_count:
                removed += 1
            continue

        title = p1 or p2 or f"Training Night {slot_date}"
        desc_lines = []
        if p1:
            desc_lines.append(f"First period: {p1}")
        if p2:
            desc_lines.append(f"Second period: {p2}")
        if uniform:
            desc_lines.append(f"Uniform: {uniform}")
        description = "\n".join(desc_lines)

        start_iso = f"{slot_date}T19:00:00+00:00"
        end_iso = f"{slot_date}T21:30:00+00:00"
        now = now_iso()
        existing = await db.events.find_one({"training_slot_id": slot_id, "source": "training_plan"}, {"_id": 0})
        event_doc = {
            "title": title,
            "description": description,
            "location": "Squadron HQ",
            "start": start_iso,
            "end": end_iso,
            "capacity": 0,
            "event_type": "standard",
            "participation": "attend",
            "points_value": 10,
            "updated_at": now,
            "updated_by": staff["id"],
            "training_slot_id": slot_id,
            "source": "training_plan",
        }
        if existing:
            await db.events.update_one({"id": existing["id"]}, {"$set": event_doc})
            updated += 1
        else:
            await db.events.insert_one({
                "id": str(uuid.uuid4()),
                **event_doc,
                "bids": [],
                "attendees": [],
                "created_at": now,
                "created_by": staff["id"],
            })
            created += 1

    return {"month": month, "created": created, "updated": updated, "removed": removed}


@api_router.post("/training-plan/{slot_id}/bid")
async def submit_training_bid(slot_id: str, payload: TrainingPlanBidCreate,
                              cfav: dict = Depends(require_roles("cfav"))):
    slot = await db.training_plan.find_one({"id": slot_id}, {"_id": 0})
    if not slot:
        raise HTTPException(status_code=404, detail="Training slot not found")
    if slot.get("no_parade"):
        raise HTTPException(status_code=400, detail="Cannot bid on a bank-holiday/no-parade slot")

    today = date.today()
    sd = date.fromisoformat(slot["slot_date"])
    month_idx = (sd.year - today.year) * 12 + (sd.month - today.month)
    if month_idx not in (1, 2, 3):
        raise HTTPException(status_code=400, detail="Bids are allowed only for months 1 to 3 ahead")

    doc = {
        "id": str(uuid.uuid4()),
        "slot_id": slot_id,
        "slot_date": slot["slot_date"],
        "cfav_id": cfav["id"],
        "cfav_name": f"{cfav.get('first_name','')} {cfav.get('last_name','')}".strip() or "CFAV",
        "is_uniformed": bool(cfav.get("is_uniformed", False)),
        "title": payload.title,
        "summary": payload.summary,
        "created_at": now_iso(),
        "status": "suggested",
    }
    await db.training_plan_bids.insert_one(doc)
    return {k: v for k, v in doc.items() if k != "_id"}


@api_router.get("/training-plan/bids")
async def list_training_bids(staff: dict = Depends(require_privileged_staff), month: Optional[str] = None):
    q = {}
    if month and re.match(r"^\d{4}-\d{2}$", month):
        q["slot_date"] = {"$gte": f"{month}-01", "$lt": f"{month}-32"}
    rows = await db.training_plan_bids.find(q, {"_id": 0}).sort("created_at", -1).to_list(3000)
    return rows


@api_router.post("/training-plan/{slot_id}/bids/{bid_id}/accept")
async def accept_training_bid(slot_id: str, bid_id: str, payload: TrainingPlanBidAccept,
                              staff: dict = Depends(require_privileged_staff)):
    bid = await db.training_plan_bids.find_one({"id": bid_id, "slot_id": slot_id}, {"_id": 0})
    if not bid:
        raise HTTPException(status_code=404, detail="Bid not found for this slot")
    slot = await db.training_plan.find_one({"id": slot_id}, {"_id": 0})
    if not slot:
        raise HTTPException(status_code=404, detail="Training slot not found")
    period = "second" if payload.period == "second" else "first"
    field = "second_period_activity" if period == "second" else "first_period_activity"
    await db.training_plan.update_one({"id": slot_id}, {"$set": {
        field: bid.get("title", ""),
        "updated_at": now_iso(),
        "updated_by": staff["id"],
    }})
    await db.training_plan_bids.update_many({"slot_id": slot_id}, {"$set": {"status": "reviewed"}})
    await db.training_plan_bids.update_one({"id": bid_id}, {"$set": {"status": "accepted", "accepted_at": now_iso(), "accepted_by": staff["id"], "accepted_period": period}})
    return {"accepted": True, "period": period}


@api_router.get("/training-plan/templates")
async def list_training_templates(staff: dict = Depends(require_privileged_staff)):
    return await db.training_plan_templates.find({}, {"_id": 0}).sort("month", -1).to_list(200)


@api_router.post("/training-plan/templates/{month}/save")
async def save_training_template(month: str, staff: dict = Depends(require_privileged_staff)):
    if not re.match(r"^\d{4}-\d{2}$", month or ""):
        raise HTTPException(status_code=400, detail="Month must be YYYY-MM")
    rows = await db.training_plan.find({
        "slot_date": {"$gte": f"{month}-01", "$lt": f"{month}-32"}
    }, {"_id": 0}).sort("slot_date", 1).to_list(500)
    tpl_rows = [{
        "slot_date": r["slot_date"],
        "first_period_activity": r.get("first_period_activity", ""),
        "second_period_activity": r.get("second_period_activity", ""),
        "uniform_needed": r.get("uniform_needed", ""),
    } for r in rows]
    await db.training_plan_templates.update_one(
        {"month": month},
        {"$set": {"month": month, "rows": tpl_rows, "updated_at": now_iso(), "updated_by": staff["id"]}},
        upsert=True,
    )
    return {"saved": True, "month": month, "rows": len(tpl_rows)}


@api_router.post("/training-plan/templates/{month}/apply")
async def apply_training_template(month: str, staff: dict = Depends(require_privileged_staff)):
    if not re.match(r"^\d{4}-\d{2}$", month or ""):
        raise HTTPException(status_code=400, detail="Month must be YYYY-MM")
    tpl = await db.training_plan_templates.find_one({"month": month}, {"_id": 0})
    if not tpl:
        raise HTTPException(status_code=404, detail="Template not found for this month")
    holidays = _uk_bank_holidays_england_wales(int(month[:4]))
    updated = 0
    for r in tpl.get("rows", []):
        sd = r.get("slot_date", "")
        if not sd.startswith(f"{month}-"):
            continue
        d = date.fromisoformat(sd)
        if d.weekday() not in (0, 3):
            continue
        is_holiday = d in holidays
        set_fields = {
            "updated_at": now_iso(),
            "updated_by": staff["id"],
            "no_parade": bool(is_holiday),
            "no_parade_reason": "Bank holiday" if is_holiday else "",
        }
        if is_holiday:
            set_fields.update({
                "first_period_activity": "NO PARADE",
                "second_period_activity": "Bank holiday",
                "uniform_needed": "N/A",
            })
        else:
            set_fields.update({
                "first_period_activity": r.get("first_period_activity", ""),
                "second_period_activity": r.get("second_period_activity", ""),
                "uniform_needed": r.get("uniform_needed", ""),
            })
        await db.training_plan.update_one({"slot_date": sd}, {"$set": set_fields}, upsert=True)
        updated += 1
    return {"applied": True, "month": month, "updated": updated}


def _escape_pdf_text(v: str) -> str:
    return (v or "").replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _simple_training_plan_pdf(month: str, rows: List[dict], role_names: dict) -> bytes:
    """Fallback PDF builder used when ReportLab is unavailable.

    Produces a valid single-page PDF with monospaced text rows.
    """
    lines = [
        "1471 Horwich Squadron RAF Air Cadets",
        f"Training Plan - {month}",
        "",
        f"Training Officer: {role_names.get('training_officer', 'TBC')}    Adjutant: {role_names.get('adjutant', 'TBC')}",
        f"Stores: {role_names.get('stores_officer', 'TBC')}    Community: {role_names.get('community_officer', 'TBC')}    H&S: {role_names.get('health_safety_officer', 'TBC')}",
        f"Generated: {datetime.now().strftime('%d %b %Y %H:%M')} UTC",
        "",
        "Day/Date         First period activity              Second period activity             Uniform",
        "----------------------------------------------------------------------------------------------",
    ]

    for r in rows[:55]:
        sd = date.fromisoformat(r["slot_date"])
        daydate = f"{sd.strftime('%a')} {sd.strftime('%d %b')}"
        p1 = (r.get("first_period_activity", "") or "")[:32]
        p2 = (r.get("second_period_activity", "") or "")[:32]
        un = (r.get("uniform_needed", "") or "")[:14]
        lines.append(f"{daydate:<16}{p1:<36}{p2:<36}{un:<14}")

    content_lines = ["BT", "/F1 10 Tf", "50 790 Td", "12 TL"]
    for idx, line in enumerate(lines):
        esc = _escape_pdf_text(line)
        if idx == 0:
            content_lines.append(f"({esc}) Tj")
        else:
            content_lines.append("T*")
            content_lines.append(f"({esc}) Tj")
    content_lines.append("ET")
    stream = "\n".join(content_lines).encode("latin-1", errors="replace")

    objs = []

    def add_obj(payload: bytes) -> int:
        objs.append(payload)
        return len(objs)

    catalog_id = add_obj(b"<< /Type /Catalog /Pages 2 0 R >>")
    pages_id = add_obj(b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>")
    page_id = add_obj(b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>")
    font_id = add_obj(b"<< /Type /Font /Subtype /Type1 /BaseFont /Courier >>")
    content_id = add_obj(f"<< /Length {len(stream)} >>\nstream\n".encode("ascii") + stream + b"\nendstream")

    out = io.BytesIO()
    out.write(b"%PDF-1.4\n")
    offsets = [0]
    for i, payload in enumerate(objs, start=1):
        offsets.append(out.tell())
        out.write(f"{i} 0 obj\n".encode("ascii"))
        out.write(payload)
        out.write(b"\nendobj\n")

    xref_pos = out.tell()
    out.write(f"xref\n0 {len(objs) + 1}\n".encode("ascii"))
    out.write(b"0000000000 65535 f \n")
    for off in offsets[1:]:
        out.write(f"{off:010d} 00000 n \n".encode("ascii"))

    out.write(
        f"trailer\n<< /Size {len(objs) + 1} /Root {catalog_id} 0 R >>\nstartxref\n{xref_pos}\n%%EOF".encode("ascii")
    )
    return out.getvalue()


def _training_plan_html_document(month: str, rows: List[dict], role_names: dict) -> str:
        generated = datetime.now().strftime("%d %b %Y %H:%M UTC")
        safe_month = html.escape(month)
        crest_url = html.escape(
                os.environ.get(
                        "SQUADRON_CREST_URL",
                        "https://customer-assets.emergentagent.com/job_283d297f-7217-4e9a-b0b5-b0baa4b4d8bf/artifacts/nmvg3tzu_1471%20crest%20transparent.png",
                )
        )

        def cell(v: str) -> str:
                return html.escape(v or "")

        def render_row(r: dict) -> str:
                sd = date.fromisoformat(r["slot_date"])
                day_name = sd.strftime("%A")
                day_date = sd.strftime("%d %b %Y")
                first = r.get("first_period_activity", "") or ""
                second = r.get("second_period_activity", "") or ""
                uniform = r.get("uniform_needed", "") or ""
                no_parade = bool(r.get("no_parade", False))
                reason = r.get("no_parade_reason", "") or ""
                cls = "no-parade" if no_parade else ""
                status = "No parade" if no_parade else "Planned"
                if no_parade and reason:
                        status = f"No parade - {reason}"
                return f"""
                <tr class=\"{cls}\">
                    <td><div class=\"day\">{cell(day_name)}</div><div class=\"date\">{cell(day_date)}</div></td>
                    <td>{cell(first)}</td>
                    <td>{cell(second)}</td>
                    <td>{cell(uniform)}</td>
                    <td><span class=\"status {'status-np' if no_parade else 'status-ok'}\">{cell(status)}</span></td>
                </tr>
                """.strip()

        rows_per_page = 14
        chunks = [rows[i:i + rows_per_page] for i in range(0, len(rows), rows_per_page)] or [[]]
        pages_html = []
        for idx, chunk in enumerate(chunks):
                page_rows = "\n".join(render_row(r) for r in chunk) if chunk else (
                        "<tr><td colspan=\"5\" class=\"empty\">No training slots were found for this month.</td></tr>"
                )
                pages_html.append(f"""
                <section class=\"page {'last-page' if idx == len(chunks) - 1 else ''}\">
                    <div class=\"sheet\">
                        <header class=\"hero\">
                            <div class=\"hero-left\">
                                <img src=\"{crest_url}\" alt=\"1471 Horwich Squadron crest\" class=\"crest\" />
                                <div>
                                    <div class=\"kicker\">1471 Horwich Squadron RAF Air Cadets</div>
                                    <h1>Training Programme</h1>
                                    <p class=\"subtitle\">Month: {safe_month} · Page {idx + 1} of {len(chunks)}</p>
                                </div>
                            </div>
                        </header>

                        <section class=\"meta\">
                            <div class=\"item\"><strong>Generated:</strong> {cell(generated)}</div>
                            <div class=\"item\"><strong>Document:</strong> Monthly training programme planner</div>
                        </section>

                        <section class=\"appointments\">
                            <div class=\"app\"><div class=\"label\">Training Officer</div><div class=\"value\">{cell(role_names.get('training_officer', 'TBC'))}</div></div>
                            <div class=\"app\"><div class=\"label\">Adjutant</div><div class=\"value\">{cell(role_names.get('adjutant', 'TBC'))}</div></div>
                            <div class=\"app\"><div class=\"label\">Stores Officer</div><div class=\"value\">{cell(role_names.get('stores_officer', 'TBC'))}</div></div>
                            <div class=\"app\"><div class=\"label\">Community Officer</div><div class=\"value\">{cell(role_names.get('community_officer', 'TBC'))}</div></div>
                            <div class=\"app\"><div class=\"label\">Health & Safety</div><div class=\"value\">{cell(role_names.get('health_safety_officer', 'TBC'))}</div></div>
                        </section>

                        <div class=\"wrap\">
                            <table>
                                <thead>
                                    <tr>
                                        <th style=\"width:17%\">Day / Date</th>
                                        <th style=\"width:28%\">First Period Activity</th>
                                        <th style=\"width:28%\">Second Period Activity</th>
                                        <th style=\"width:14%\">Uniform Needed</th>
                                        <th style=\"width:13%\">Status</th>
                                    </tr>
                                </thead>
                                <tbody>
                                    {page_rows}
                                </tbody>
                            </table>
                        </div>

                        <footer class=\"foot\">
                            1471 Horwich Squadron RAF Air Cadets · Internal planning copy · For authorised staff use.
                        </footer>
                    </div>
                </section>
                """)

        return f"""<!doctype html>
<html lang=\"en\">
<head>
    <meta charset=\"utf-8\" />
    <meta name=\"viewport\" content=\"width=device-width, initial-scale=1\" />
    <title>1471 Horwich Squadron - Training Plan {safe_month}</title>
    <style>
        :root {{
            --raf-navy: #071A2F;
            --raf-blue: #00529B;
            --raf-sky: #DDEFFC;
            --raf-red: #C60C30;
            --ink: #1B2430;
            --slate: #5B6B78;
            --paper: #F5FAFF;
            --border: #C8D9E6;
        }}
        * {{ box-sizing: border-box; }}
        body {{
            margin: 0;
            font-family: "Segoe UI", Tahoma, Arial, sans-serif;
            color: var(--ink);
            background: linear-gradient(165deg, #f7fbff 0%, #edf5fb 100%);
            padding: 28px;
        }}
        .page {{ margin-bottom: 24px; }}
        .page.last-page {{ margin-bottom: 0; }}
        .sheet {{
            max-width: 1100px;
            margin: 0 auto;
            background: #fff;
            border: 1px solid var(--border);
            box-shadow: 0 10px 30px rgba(7, 26, 47, 0.08);
            overflow: hidden;
        }}
        .hero {{
            background: linear-gradient(135deg, var(--raf-navy) 0%, var(--raf-blue) 100%);
            color: #fff;
            padding: 22px 26px 20px;
            border-bottom: 6px solid var(--raf-red);
            position: relative;
        }}
        .hero-left {{ display: flex; align-items: center; gap: 14px; position: relative; z-index: 1; }}
        .crest {{ width: 64px; height: 64px; object-fit: contain; filter: drop-shadow(0 4px 10px rgba(0,0,0,0.25)); flex: 0 0 auto; }}
        .hero::after {{
            content: "";
            position: absolute;
            right: -80px;
            top: -80px;
            width: 210px;
            height: 210px;
            border-radius: 50%;
            background: radial-gradient(circle at 30% 30%, rgba(255,255,255,0.16), rgba(255,255,255,0.03));
        }}
        .kicker {{ letter-spacing: .12em; text-transform: uppercase; font-size: 11px; color: #9ed1ff; font-weight: 700; }}
        h1 {{ margin: 8px 0 4px; font-size: 30px; line-height: 1.08; }}
        .subtitle {{ margin: 0; color: #dcecff; font-size: 14px; }}
        .meta {{
            display: grid;
            grid-template-columns: repeat(2, minmax(0, 1fr));
            gap: 8px 12px;
            padding: 16px 20px;
            background: var(--paper);
            border-bottom: 1px solid var(--border);
            font-size: 13px;
        }}
        .meta .item strong {{ color: var(--raf-navy); }}
        .appointments {{
            display: grid;
            grid-template-columns: repeat(5, minmax(0, 1fr));
            gap: 0;
            border-bottom: 1px solid var(--border);
            background: #fff;
        }}
        .app {{ padding: 12px 14px; border-right: 1px solid var(--border); min-height: 62px; }}
        .app:last-child {{ border-right: 0; }}
        .app .label {{ font-size: 10px; text-transform: uppercase; letter-spacing: .08em; color: var(--slate); font-weight: 700; }}
        .app .value {{ margin-top: 5px; font-size: 13px; color: var(--raf-navy); font-weight: 700; }}
        .wrap {{ padding: 16px 18px 20px; }}
        table {{ width: 100%; border-collapse: collapse; font-size: 13px; }}
        th {{
            background: var(--raf-navy);
            color: #fff;
            text-align: left;
            font-size: 11px;
            letter-spacing: .06em;
            text-transform: uppercase;
            padding: 10px 9px;
            border: 1px solid #122944;
        }}
        td {{ border: 1px solid var(--border); padding: 9px; vertical-align: top; }}
        td .day {{ font-weight: 700; color: var(--raf-navy); }}
        td .date {{ font-size: 11px; color: var(--slate); margin-top: 2px; }}
        tr:nth-child(even) td {{ background: #fcfeff; }}
        tr.no-parade td {{ background: #fff3f3; }}
        .status {{ display: inline-block; border-radius: 999px; padding: 3px 8px; font-size: 11px; font-weight: 700; }}
        .status-ok {{ background: #e7f6ec; color: #116c2f; }}
        .status-np {{ background: #fde8eb; color: #a50f2f; }}
        .empty {{ text-align: center; color: var(--slate); padding: 20px 8px; }}
        .foot {{
            border-top: 1px solid var(--border);
            padding: 12px 18px 16px;
            font-size: 11px;
            color: var(--slate);
            background: #fbfdff;
        }}
        @media print {{
            @page {{ size: A4 portrait; margin: 10mm; }}
            body {{ background: #fff; padding: 0; }}
            .page {{ page-break-after: always; break-after: page; margin: 0; }}
            .page.last-page {{ page-break-after: auto; break-after: auto; }}
            .sheet {{ border: 0; box-shadow: none; max-width: none; }}
            .hero {{ padding-top: 14px; padding-bottom: 14px; }}
            h1 {{ font-size: 24px; }}
            .subtitle {{ font-size: 12px; }}
            .meta {{ padding-top: 10px; padding-bottom: 10px; }}
            .appointments {{ break-inside: avoid; }}
            table {{ break-inside: auto; }}
            tr {{ break-inside: avoid; }}
        }}
    </style>
</head>
<body>
    {''.join(pages_html)}
</body>
</html>
"""


@api_router.get("/training-plan/a4")
async def training_plan_a4(month: str, format: Optional[str] = None, staff: dict = Depends(require_privileged_staff)):
    if not re.match(r"^\d{4}-\d{2}$", month or ""):
        raise HTTPException(status_code=400, detail="Month must be YYYY-MM")
    rows = await db.training_plan.find({
        "slot_date": {"$gte": f"{month}-01", "$lt": f"{month}-32"}
    }, {"_id": 0}).sort("slot_date", 1).to_list(500)

    app_doc = await _appointments_doc()
    v = app_doc.get("value", {})
    role_names = {}
    for k in APPOINTMENT_KEYS:
        uid = v.get(k)
        if not uid:
            role_names[k] = "TBC"
            continue
        u = await db.users.find_one({"id": uid}, {"_id": 0, "first_name": 1, "last_name": 1})
        role_names[k] = (f"{u.get('first_name','')} {u.get('last_name','')}".strip() if u else "TBC")

    if (format or "").lower() == "html":
        html_doc = _training_plan_html_document(month, rows, role_names)
        return Response(
            content=html_doc,
            media_type="text/html; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="training-plan-{month}.html"'},
        )

    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
    except Exception:
        pdf = _simple_training_plan_pdf(month, rows, role_names)
        return Response(
            content=pdf,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="training-plan-{month}.pdf"'},
        )

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    w, h = A4

    y = h - 32
    c.setFillColorRGB(0.0, 0.18, 0.37)
    c.rect(0, h - 70, w, 70, fill=1, stroke=0)
    c.setFillColorRGB(0.78, 0.05, 0.19)
    c.rect(0, h - 75, w, 5, fill=1, stroke=0)
    c.setFillColorRGB(1, 1, 1)
    c.setFont("Helvetica-Bold", 16)
    c.drawString(40, y, "1471 Horwich Squadron RAF Air Cadets")
    c.setFont("Helvetica", 11)
    c.drawString(40, y - 18, f"Training Plan - {month}")

    y = h - 90
    c.setFillColorRGB(0, 0, 0)
    c.setFont("Helvetica", 8)
    c.drawString(40, y, f"Training Officer: {role_names['training_officer']}")
    c.drawString(220, y, f"Adjutant: {role_names['adjutant']}")
    c.drawString(360, y, f"Stores: {role_names['stores_officer']}")
    c.drawString(480, y, f"H&S: {role_names['health_safety_officer']}")
    y -= 14
    c.drawString(40, y, f"Community: {role_names['community_officer']}")
    c.drawString(220, y, f"Generated: {datetime.now().strftime('%d %b %Y %H:%M')} UTC")
    y -= 18
    c.setFont("Helvetica", 9)
    c.drawString(40, y, "Day/Date")
    c.drawString(130, y, "First period activity")
    c.drawString(310, y, "Second period activity")
    c.drawString(485, y, "Uniform needed")
    y -= 8
    c.line(40, y, w - 40, y)
    y -= 14

    for r in rows:
        if y < 60:
            c.showPage()
            y = h - 40
            c.setFont("Helvetica", 9)
        sd = date.fromisoformat(r["slot_date"])
        daydate = f"{sd.strftime('%a')} {sd.strftime('%d %b')}"
        p1 = r.get("first_period_activity", "")
        p2 = r.get("second_period_activity", "")
        un = r.get("uniform_needed", "")
        c.drawString(40, y, daydate[:16])
        c.drawString(130, y, p1[:34])
        c.drawString(310, y, p2[:32])
        c.drawString(485, y, un[:16])
        y -= 14

    c.save()
    pdf = buf.getvalue()
    return Response(
        content=pdf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="training-plan-{month}.pdf"'},
    )


# ---------------------------------------------------------------------------
# Events
# ---------------------------------------------------------------------------
def event_view(e: dict, user_id: str, staff: bool) -> dict:
    bids = e.get("bids", [])
    attendees = e.get("attendees", [])
    cap = int(e.get("capacity", 0))
    bid_count = len(bids)
    if cap <= 0:
        colour = "green"
    else:
        ratio = bid_count / cap
        colour = "green" if ratio < 0.5 else ("amber" if ratio < 1.0 else "red")
    view = {
        "id": e["id"], "title": e["title"], "description": e.get("description", ""),
        "location": e.get("location", ""), "link_url": e.get("link_url", ""),
        "attachment_ids": e.get("attachment_ids", []),
        "start": e["start"], "end": e.get("end"),
        "capacity": cap, "event_type": e.get("event_type", "standard"),
        "participation": e.get("participation", "attend"),
        "points_value": int(e.get("points_value", 0)),
        "bid_count": bid_count, "attendee_count": len(attendees),
        "colour": colour, "my_bid": user_id in bids, "my_attendance": user_id in attendees,
    }
    if staff:
        view["bids"] = bids
        view["attendees"] = attendees
    return view


@api_router.get("/events")
async def list_events(user: dict = Depends(get_current_user)):
    rows = await db.events.find({}, {"_id": 0}).sort("start", 1).to_list(3000)
    staff = user["role"] in STAFF_ROLES
    return [event_view(e, user["id"], staff) for e in rows]


@api_router.post("/events")
async def create_event(payload: EventCreate, staff: dict = Depends(require_staff)):
    e = {
        "id": str(uuid.uuid4()),
        "title": payload.title,
        "description": payload.description,
        "location": payload.location,
        "link_url": (payload.link_url or DEFAULT_EVENT_LINK).strip(),
        "attachment_ids": payload.attachment_ids or [],
        "start": payload.start,
        "end": payload.end,
        "capacity": int(payload.capacity or 0),
        "event_type": payload.event_type or "standard",
        "participation": payload.participation or "attend",
        "points_value": int(payload.points_value or 0),
        "bids": [],
        "attendees": [],
        "created_at": now_iso(),
        "updated_at": now_iso(),
        "created_by": staff["id"],
    }
    await db.events.insert_one(e)
    return event_view(e, staff["id"], True)


@api_router.patch("/events/{event_id}")
async def update_event(event_id: str, payload: EventUpdate, staff: dict = Depends(require_staff)):
    updates = {k: v for k, v in payload.model_dump().items() if v is not None}
    if not updates:
        raise HTTPException(status_code=400, detail="No fields to update")
    if "capacity" in updates:
        updates["capacity"] = int(updates["capacity"] or 0)
    if "points_value" in updates:
        updates["points_value"] = int(updates["points_value"] or 0)
    updates["updated_at"] = now_iso()
    e = await db.events.find_one_and_update(
        {"id": event_id}, {"$set": updates}, projection={"_id": 0}, return_document=True)
    if not e:
        raise HTTPException(status_code=404, detail="Event not found")
    return event_view(e, staff["id"], True)


@api_router.get("/events/{event_id}")
async def get_event(event_id: str, user: dict = Depends(get_current_user)):
    e = await db.events.find_one({"id": event_id}, {"_id": 0})
    if not e:
        raise HTTPException(status_code=404, detail="Event not found")
    staff = user["role"] in STAFF_ROLES
    return await _event_detail_view(e, user["id"], staff)


@api_router.delete("/events/{event_id}")
async def delete_event(event_id: str, staff: dict = Depends(require_staff)):
    res = await db.events.delete_one({"id": event_id})
    if res.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Event not found")
    return {"deleted": True}


@api_router.post("/events/{event_id}/bid")
async def toggle_bid(event_id: str, user: dict = Depends(require_roles("cadet"))):
    e = await db.events.find_one({"id": event_id})
    if not e:
        raise HTTPException(status_code=404, detail="Event not found")
    bids = e.get("bids", [])
    if user["id"] in bids:
        bids.remove(user["id"])
        action = "withdrawn"
    else:
        cap = int(e.get("capacity", 0))
        if cap > 0 and len(bids) >= cap:
            raise HTTPException(status_code=400, detail="This event is already full")
        bids.append(user["id"])
        action = "placed"
    await db.events.update_one({"id": event_id}, {"$set": {"bids": bids}})
    return {"action": action, "bid_count": len(bids), "my_bid": user["id"] in bids}


@api_router.post("/events/{event_id}/attendance")
async def set_attendance(event_id: str, payload: AttendanceUpdate, staff: dict = Depends(require_staff)):
    e = await db.events.find_one({"id": event_id})
    if not e:
        raise HTTPException(status_code=404, detail="Event not found")
    await db.events.update_one({"id": event_id}, {"$set": {"attendees": payload.attendee_ids}})
    e["attendees"] = payload.attendee_ids
    return event_view(e, staff["id"], True)


# ---------------------------------------------------------------------------
# Notices
# ---------------------------------------------------------------------------
@api_router.get("/notices")
async def list_notices(user: dict = Depends(get_current_user)):
    if user["role"] in STAFF_ROLES:
        notices = await db.notices.find({}, {"_id": 0}).sort("created_at", -1).to_list(1000)
    else:
        notices = await db.notices.find({"roles": user["role"]}, {"_id": 0}).sort("created_at", -1).to_list(1000)
    acks = await db.notice_acks.find({"user_id": user["id"]}, {"_id": 0, "notice_id": 1}).to_list(2000)
    acked = {a["notice_id"] for a in acks}
    for n in notices:
        n["acknowledged"] = n["id"] in acked
    return notices


@api_router.get("/notices/pending")
async def pending_notices(user: dict = Depends(get_current_user)):
    notices = await db.notices.find(
        {"roles": user["role"], "requires_ack": True}, {"_id": 0}).sort("created_at", 1).to_list(1000)
    acks = await db.notice_acks.find({"user_id": user["id"]}, {"_id": 0, "notice_id": 1}).to_list(2000)
    acked = {a["notice_id"] for a in acks}
    return [n for n in notices if n["id"] not in acked]


@api_router.post("/notices/{notice_id}/ack")
async def ack_notice(notice_id: str, user: dict = Depends(get_current_user)):
    await db.notice_acks.update_one(
        {"notice_id": notice_id, "user_id": user["id"]},
        {"$set": {"notice_id": notice_id, "user_id": user["id"], "at": now_iso()}},
        upsert=True)
    return {"acknowledged": True}


@api_router.post("/notices")
async def create_notice(payload: NoticeCreate, staff: dict = Depends(require_staff)):
    roles = [r for r in payload.roles if r in ALL_ROLES]
    if not roles:
        raise HTTPException(status_code=400, detail="Select at least one valid audience")
    notice = {"id": str(uuid.uuid4()), "title": payload.title, "body": payload.body,
              "roles": roles, "requires_ack": payload.requires_ack,
              "created_by": f"{staff.get('first_name','')} {staff.get('last_name','')}".strip() or "Staff",
              "created_at": now_iso()}
    await db.notices.insert_one(notice)
    notice.pop("_id", None)
    recipients = await db.users.find({"role": {"$in": roles}}, {"id": 1}).to_list(5000)
    for u in recipients:
        await push_to_user(u["id"], f"Notice: {payload.title}", payload.body, "/portal")
    return {k: v for k, v in notice.items() if k != "_id"}


@api_router.delete("/notices/{notice_id}")
async def delete_notice(notice_id: str, staff: dict = Depends(require_staff)):
    res = await db.notices.delete_one({"id": notice_id})
    if res.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Notice not found")
    await db.notice_acks.delete_many({"notice_id": notice_id})
    return {"deleted": True}


# ---------------------------------------------------------------------------
# Messages (member <-> staff board)
# ---------------------------------------------------------------------------
@api_router.get("/messages/thread")
async def my_thread(user: dict = Depends(require_roles("cadet", "parent"))):
    msgs = await db.messages.find({"member_id": user["id"]}, {"_id": 0}).sort("created_at", 1).to_list(1000)
    await db.messages.update_many(
        {"member_id": user["id"], "from_staff": True, "read_by_member": {"$ne": True}},
        {"$set": {"read_by_member": True}})
    return msgs


@api_router.post("/messages/thread")
async def post_to_thread(payload: MessageCreate, user: dict = Depends(require_roles("cadet", "parent"))):
    msg = {"id": str(uuid.uuid4()), "member_id": user["id"],
           "member_name": f"{user.get('first_name','')} {user.get('last_name','')}".strip(),
           "from_staff": False, "author": user.get("first_name", "Member"),
           "body": payload.body, "created_at": now_iso(),
           "read_by_member": True, "read_by_staff": False}
    await db.messages.insert_one(msg)
    return {k: v for k, v in msg.items() if k != "_id"}


@api_router.get("/messages/threads")
async def staff_threads(staff: dict = Depends(require_staff)):
    pipeline = [
        {"$sort": {"created_at": -1}},
        {"$group": {"_id": "$member_id",
                    "member_name": {"$first": "$member_name"},
                    "last_body": {"$first": "$body"},
                    "last_at": {"$first": "$created_at"},
                    "unread": {"$sum": {"$cond": [{"$and": [
                        {"$eq": ["$from_staff", False]},
                        {"$ne": ["$read_by_staff", True]}]}, 1, 0]}}}},
        {"$sort": {"last_at": -1}},
    ]
    rows = await db.messages.aggregate(pipeline).to_list(1000)
    out = []
    for r in rows:
        member = await db.users.find_one({"id": r["_id"]}, {"_id": 0, "password_hash": 0})
        out.append({"member_id": r["_id"],
                    "member_name": r.get("member_name") or (member and f"{member.get('first_name','')} {member.get('last_name','')}".strip()) or "Member",
                    "member_role": member["role"] if member else "member",
                    "last_body": r["last_body"], "last_at": r["last_at"], "unread": r["unread"]})
    return out


@api_router.get("/messages/thread/{member_id}")
async def staff_view_thread(member_id: str, staff: dict = Depends(require_staff)):
    msgs = await db.messages.find({"member_id": member_id}, {"_id": 0}).sort("created_at", 1).to_list(1000)
    await db.messages.update_many(
        {"member_id": member_id, "from_staff": False, "read_by_staff": {"$ne": True}},
        {"$set": {"read_by_staff": True}})
    return msgs


@api_router.post("/messages/thread/{member_id}")
async def staff_reply(member_id: str, payload: MessageCreate, staff: dict = Depends(require_staff)):
    member = await db.users.find_one({"id": member_id}, {"_id": 0, "password_hash": 0})
    if not member:
        raise HTTPException(status_code=404, detail="Member not found")
    msg = {"id": str(uuid.uuid4()), "member_id": member_id,
           "member_name": f"{member.get('first_name','')} {member.get('last_name','')}".strip(),
           "from_staff": True, "author": f"{staff.get('first_name','')} {staff.get('last_name','')}".strip() or "Staff",
           "body": payload.body, "created_at": now_iso(),
           "read_by_member": False, "read_by_staff": True}
    await db.messages.insert_one(msg)
    return {k: v for k, v in msg.items() if k != "_id"}


# ---------------------------------------------------------------------------
# SMS support (CFAV reset request -> admin reply in dashboard)
# ---------------------------------------------------------------------------
def _sms_reset_email_html(requester_name: str, requester_email: str, note: str = "") -> str:
    note_html = f'<p style="padding:12px;background:#EAF5F8;border-left:3px solid #002F5F;">{note}</p>' if note else ""
    inner = f"""
      <p><strong>SMS password reset request</strong></p>
      <p><strong>From:</strong> {requester_name}</p>
      <p><strong>Email:</strong> {requester_email}</p>
      {note_html}
      <p>Please action this SMS password reset. Replies should be sent in the portal SMS support thread rather than by email.</p>
    """
    return _email_shell("SMS password reset request", inner)


def _sms_thread_out(t: dict) -> dict:
    return {
        "id": t["id"],
        "requester_id": t["requester_id"],
        "requester_name": t.get("requester_name", "CFAV"),
        "requester_email": t.get("requester_email", ""),
        "status": t.get("status", "open"),
        "created_at": t.get("created_at"),
        "updated_at": t.get("updated_at"),
        "subject": t.get("subject", "SMS password reset"),
        "email_status": t.get("email_status", "unknown"),
        "messages": t.get("messages", []),
    }


@api_router.post("/sms-support/request")
async def sms_support_request(payload: SmsResetRequestCreate, user: dict = Depends(require_roles("cfav"))):
    requester_name = f"{user.get('first_name','')} {user.get('last_name','')}".strip() or "CFAV"
    note = (payload.note or "").strip()
    initial = note or "Please reset my SMS password."
    msg = {
        "id": str(uuid.uuid4()),
        "author_id": user["id"],
        "author_name": requester_name,
        "from_admin": False,
        "body": initial,
        "created_at": now_iso(),
        "read_by_requester": True,
        "read_by_admin": False,
    }
    thread = {
        "id": str(uuid.uuid4()),
        "requester_id": user["id"],
        "requester_name": requester_name,
        "requester_email": user.get("email", ""),
        "status": "open",
        "subject": "SMS password reset",
        "created_at": now_iso(),
        "updated_at": now_iso(),
        "email_status": "pending",
        "messages": [msg],
    }
    email_status = await send_email(
        SMS_RESET_EMAIL,
        f"SMS password reset request - {requester_name}",
        _sms_reset_email_html(requester_name, user.get("email", ""), note),
        reply_to=user.get("email", "") or None,
        from_email=mailbox("admin"),
    )
    thread["email_status"] = email_status
    await db.sms_support_threads.insert_one(thread)

    admins = await db.users.find({"role": "admin"}, {"_id": 0, "id": 1}).to_list(100)
    notice_docs = [{
        "id": str(uuid.uuid4()),
        "user_id": a["id"],
        "title": "SMS reset request",
        "body": f"{requester_name} requested an SMS password reset.",
        "from_name": requester_name,
        "kind": "sms_support",
        "channels": ["dashboard"],
        "read": False,
        "created_at": now_iso(),
    } for a in admins]
    if notice_docs:
        await db.notifications.insert_many(notice_docs)
        for n in notice_docs:
            await push_to_user(n["user_id"], n["title"], n["body"], "/portal")
    return {"thread": _sms_thread_out(thread), "email_status": email_status}


@api_router.get("/sms-support/threads")
async def sms_support_threads(admin: dict = Depends(require_roles("admin"))):
    rows = await db.sms_support_threads.find({}, {"_id": 0}).sort("updated_at", -1).to_list(1000)
    out = []
    for t in rows:
        unread = sum(1 for m in t.get("messages", []) if not m.get("from_admin") and not m.get("read_by_admin", False))
        x = _sms_thread_out(t)
        x.pop("messages", None)
        x["unread"] = unread
        x["last_body"] = (t.get("messages") or [{}])[-1].get("body", "")
        out.append(x)
    return out


@api_router.get("/sms-support/my-threads")
async def sms_support_my_threads(user: dict = Depends(require_roles("cfav"))):
    rows = await db.sms_support_threads.find({"requester_id": user["id"]}, {"_id": 0}).sort("updated_at", -1).to_list(200)
    out = []
    for t in rows:
        unread = sum(1 for m in t.get("messages", []) if m.get("from_admin") and not m.get("read_by_requester", False))
        x = _sms_thread_out(t)
        x.pop("messages", None)
        x["unread"] = unread
        x["last_body"] = (t.get("messages") or [{}])[-1].get("body", "")
        out.append(x)
    return out


@api_router.get("/sms-support/thread/{thread_id}")
async def sms_support_view_thread(thread_id: str, user: dict = Depends(require_roles("admin", "cfav"))):
    t = await db.sms_support_threads.find_one({"id": thread_id}, {"_id": 0})
    if not t:
        raise HTTPException(status_code=404, detail="Thread not found")
    if user["role"] != "admin" and t.get("requester_id") != user["id"]:
        raise HTTPException(status_code=403, detail="Insufficient permissions")

    field = "messages.$[m].read_by_admin" if user["role"] == "admin" else "messages.$[m].read_by_requester"
    arr_filter = [{"m.from_admin": False}] if user["role"] == "admin" else [{"m.from_admin": True}]
    await db.sms_support_threads.update_one(
        {"id": thread_id},
        {"$set": {field: True}},
        array_filters=arr_filter,
    )
    t2 = await db.sms_support_threads.find_one({"id": thread_id}, {"_id": 0})
    return _sms_thread_out(t2)


@api_router.post("/sms-support/thread/{thread_id}/reply")
async def sms_support_reply(thread_id: str, payload: MessageCreate, user: dict = Depends(require_roles("admin", "cfav"))):
    t = await db.sms_support_threads.find_one({"id": thread_id}, {"_id": 0})
    if not t:
        raise HTTPException(status_code=404, detail="Thread not found")
    if user["role"] != "admin" and t.get("requester_id") != user["id"]:
        raise HTTPException(status_code=403, detail="Insufficient permissions")

    from_admin = user["role"] == "admin"
    author_name = f"{user.get('first_name','')} {user.get('last_name','')}".strip() or ("Admin" if from_admin else "CFAV")
    msg = {
        "id": str(uuid.uuid4()),
        "author_id": user["id"],
        "author_name": author_name,
        "from_admin": from_admin,
        "body": payload.body,
        "created_at": now_iso(),
        "read_by_requester": from_admin is False,
        "read_by_admin": from_admin is True,
    }
    await db.sms_support_threads.update_one(
        {"id": thread_id},
        {"$push": {"messages": msg}, "$set": {"updated_at": now_iso()}},
    )
    return msg


# ---------------------------------------------------------------------------
# Broadcast messages / notifications (staff -> targeted members)
# ---------------------------------------------------------------------------
@api_router.post("/broadcast")
async def create_broadcast(payload: BroadcastCreate, staff: dict = Depends(require_privileged_staff)):
    users = await resolve_recipients(payload.audience)
    if not users:
        raise HTTPException(status_code=400, detail="No recipients match this audience")
    from_name = f"{staff.get('first_name','')} {staff.get('last_name','')}".strip() or "Squadron Staff"
    from_email = await staff_from_email(staff)
    attachments = await _fetch_attachments(payload.attachment_ids)
    attach_note = ""
    if attachments:
        names = "\n".join(f"- {a['filename']}" for a in attachments)
        attach_note = f"\n\nAttachments:\n{names}"
    email_html = _broadcast_email_html(payload.title, payload.body, from_name, attachments, payload.base_url)
    links = [{"url": f"/api/attachments/{a['id']}/download", "label": a["filename"]} for a in attachments]
    result = await deliver_broadcast(
        users,
        payload.title,
        payload.body + attach_note,
        payload.channels,
        from_name,
        "message",
        email_html=email_html,
        link=links[0]["url"] if links else None,
        link_label=links[0]["label"] if links else None,
        links=links,
        from_email=from_email,
    )
    await db.broadcasts.insert_one({
        "id": str(uuid.uuid4()), "title": payload.title, "body": payload.body,
        "channels": result and payload.channels, "audience": payload.audience.model_dump(),
        "attachment_ids": payload.attachment_ids,
        "sent_by": from_name, "created_at": now_iso(), "result": result})
    return result


@api_router.get("/notifications")
async def my_notifications(user: dict = Depends(get_current_user)):
    return await db.notifications.find({"user_id": user["id"]}, {"_id": 0}).sort("created_at", -1).to_list(500)


@api_router.get("/notifications/unread-count")
async def unread_count(user: dict = Depends(get_current_user)):
    n = await db.notifications.count_documents({"user_id": user["id"], "read": False})
    return {"count": n}


@api_router.post("/notifications/{notification_id}/read")
async def mark_notification_read(notification_id: str, user: dict = Depends(get_current_user)):
    await db.notifications.update_one({"id": notification_id, "user_id": user["id"]},
                                      {"$set": {"read": True}})
    return {"read": True}


@api_router.post("/notifications/read-all")
async def mark_all_read(user: dict = Depends(get_current_user)):
    await db.notifications.update_many({"user_id": user["id"], "read": False},
                                       {"$set": {"read": True}})
    return {"read": True}


# ---------------------------------------------------------------------------
# Website content CMS (page text + image overrides)
# ---------------------------------------------------------------------------
def _site_doc_out(d: dict) -> dict:
    return {
        "path": d.get("path", "/"),
        "texts": d.get("texts", {}),
        "images": d.get("images", {}),
        "updated_at": d.get("updated_at"),
        "updated_by": d.get("updated_by", ""),
    }


@api_router.get("/site-content/page")
async def get_site_content_page(path: str):
    p = (path or "/").strip()
    if not p.startswith("/"):
        p = f"/{p}"
    doc = await db.site_content.find_one({"path": p}, {"_id": 0})
    if not doc:
        return {"path": p, "texts": {}, "images": {}, "updated_at": None, "updated_by": ""}
    return _site_doc_out(doc)


@api_router.get("/site-content/pages")
async def list_site_content_pages(staff: dict = Depends(require_staff)):
    rows = await db.site_content.find({}, {"_id": 0}).sort("path", 1).to_list(1000)
    return [_site_doc_out(r) for r in rows]


@api_router.put("/site-content/pages")
async def upsert_site_content_page(payload: SiteContentUpdate, staff: dict = Depends(require_staff)):
    p = payload.path.strip()
    if not p.startswith("/"):
        p = f"/{p}"
    if p.startswith("/portal"):
        raise HTTPException(status_code=400, detail="Portal routes cannot be edited with site CMS")
    images = {k: v.model_dump() for k, v in payload.images.items()}
    cleaned_texts = {k: v for k, v in payload.texts.items() if isinstance(v, str)}
    doc = {
        "path": p,
        "texts": cleaned_texts,
        "images": images,
        "updated_at": now_iso(),
        "updated_by": f"{staff.get('first_name','')} {staff.get('last_name','')}".strip() or staff.get("email", "Staff"),
    }
    await db.site_content.update_one({"path": p}, {"$set": doc}, upsert=True)
    return doc


@api_router.delete("/site-content/pages")
async def delete_site_content_page(path: str, staff: dict = Depends(require_staff)):
    p = (path or "").strip()
    if not p:
        raise HTTPException(status_code=400, detail="Path is required")
    if not p.startswith("/"):
        p = f"/{p}"
    res = await db.site_content.delete_one({"path": p})
    return {"deleted": res.deleted_count > 0}


# ---------------------------------------------------------------------------
# Newsletters (staff compose + preview + send)
# ---------------------------------------------------------------------------
def _newsletter_out(n: dict) -> dict:
    return {k: v for k, v in n.items() if k != "_id"}


@api_router.get("/newsletters")
async def list_newsletters(staff: dict = Depends(require_privileged_staff)):
    rows = await db.newsletters.find({}, {"_id": 0}).sort("created_at", -1).to_list(500)
    return rows


@api_router.post("/newsletters")
async def create_newsletter(payload: NewsletterCreate, staff: dict = Depends(require_privileged_staff)):
    nl = {"id": str(uuid.uuid4()), **payload.model_dump(), "status": "draft",
          "created_by": f"{staff.get('first_name','')} {staff.get('last_name','')}".strip() or "Staff",
          "created_at": now_iso(), "sent_at": None, "result": None}
    await db.newsletters.insert_one(nl)
    return _newsletter_out(nl)


@api_router.patch("/newsletters/{newsletter_id}")
async def update_newsletter(newsletter_id: str, payload: NewsletterCreate, staff: dict = Depends(require_privileged_staff)):
    n = await db.newsletters.find_one_and_update(
        {"id": newsletter_id}, {"$set": payload.model_dump()},
        projection={"_id": 0}, return_document=True)
    if not n:
        raise HTTPException(status_code=404, detail="Newsletter not found")
    return n


@api_router.delete("/newsletters/{newsletter_id}")
async def delete_newsletter(newsletter_id: str, staff: dict = Depends(require_privileged_staff)):
    res = await db.newsletters.delete_one({"id": newsletter_id})
    if res.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Newsletter not found")
    return {"deleted": True}


@api_router.post("/newsletters/preview")
async def preview_newsletter(payload: NewsletterCreate, staff: dict = Depends(require_privileged_staff)):
    attachments = await _fetch_attachments(payload.attachment_ids)
    return {"html": _newsletter_email_html(payload.model_dump(), attachments)}


@api_router.post("/newsletters/{newsletter_id}/send")
async def send_newsletter(newsletter_id: str, payload: NewsletterSend, staff: dict = Depends(require_privileged_staff)):
    nl = await db.newsletters.find_one({"id": newsletter_id}, {"_id": 0})
    if not nl:
        raise HTTPException(status_code=404, detail="Newsletter not found")
    users = await resolve_recipients(payload.audience)
    if not users:
        raise HTTPException(status_code=400, detail="No recipients match this audience")
    from_name = nl.get("created_by") or "Squadron Staff"
    from_email = await staff_from_email(staff)
    attachments = await _fetch_attachments(nl.get("attachment_ids") or [])
    attach_note = ""
    if attachments:
        names = "\n".join(f"- {a['filename']}" for a in attachments)
        attach_note = f"\n\nAttachments:\n{names}"
    dash_body = (nl.get("intro", "") + ("\n\n" if nl.get("intro") else "") + nl.get("body", "")).strip()
    dash_body = (dash_body + attach_note).strip()
    email_html = _newsletter_email_html(nl, attachments, payload.base_url)
    links = [{"url": f"/api/attachments/{a['id']}/download", "label": a["filename"]} for a in attachments]
    result = await deliver_broadcast(users, nl["subject"], dash_body, payload.channels,
                                     from_name, "newsletter", email_html=email_html,
                                     link=links[0]["url"] if links else None,
                                     link_label=links[0]["label"] if links else None,
                                     links=links, from_email=from_email)
    await db.newsletters.update_one({"id": newsletter_id},
                                    {"$set": {"status": "sent", "sent_at": now_iso(),
                                              "audience": payload.audience.model_dump(),
                                              "channels": payload.channels, "result": result}})
    return result


# ---------------------------------------------------------------------------
# Calendar: ICS subscribe feed + Word (.docx) training-programme import
# ---------------------------------------------------------------------------
class ImportedEvent(BaseModel):
    title: str = Field(..., min_length=1, max_length=200)
    start: str
    end: Optional[str] = None
    description: str = ""
    location: str = ""
    model_config = ConfigDict(extra="ignore")


class EventsImport(BaseModel):
    events: List[ImportedEvent]
    model_config = ConfigDict(extra="ignore")


def _ics_escape(t: str) -> str:
    return (t or "").replace("\\", "\\\\").replace(";", "\\;").replace(",", "\\,").replace("\n", "\\n")


def _ics_dt(iso: str) -> Optional[str]:
    if not iso:
        return None
    try:
        dt = datetime.fromisoformat(iso.replace("Z", "+00:00"))
    except Exception:
        return None
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=timezone.utc)
    return dt.astimezone(timezone.utc).strftime("%Y%m%dT%H%M%SZ")


@api_router.get("/calendar/events.ics")
async def events_ics():
    rows = await db.events.find({}, {"_id": 0}).sort("start", 1).to_list(2000)
    lines = ["BEGIN:VCALENDAR", "VERSION:2.0",
             "PRODID:-//1471 Horwich Squadron//RAF Air Cadets//EN", "CALSCALE:GREGORIAN",
             "METHOD:PUBLISH", "X-WR-CALNAME:1471 Horwich Squadron Events"]
    for e in rows:
        dtstart = _ics_dt(e.get("start", ""))
        if not dtstart:
            continue
        lines += ["BEGIN:VEVENT", f"UID:{e['id']}@1471horwich.org.uk",
                  f"DTSTAMP:{_ics_dt(e.get('created_at') or now_iso()) or dtstart}",
                  f"DTSTART:{dtstart}"]
        dtend = _ics_dt(e.get("end") or "")
        if dtend:
            lines.append(f"DTEND:{dtend}")
        lines.append(f"SUMMARY:{_ics_escape(e.get('title', ''))}")
        if e.get("location"):
            lines.append(f"LOCATION:{_ics_escape(e['location'])}")
        if e.get("description"):
            lines.append(f"DESCRIPTION:{_ics_escape(e['description'])}")
        lines.append("END:VEVENT")
    lines.append("END:VCALENDAR")
    body = "\r\n".join(lines) + "\r\n"
    return Response(content=body, media_type="text/calendar; charset=utf-8",
                    headers={"Content-Disposition": 'inline; filename="1471-horwich-events.ics"'})


def _parse_docx_events(data: bytes) -> list:
    from docx import Document as DocxDocument
    from dateutil import parser as dateparser
    doc = DocxDocument(io.BytesIO(data))
    rows = []
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                rows.append(cells)
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            rows.append([t])
    year = datetime.now(timezone.utc).year
    headers = {"date", "activity", "event", "events", "notes", "week", "time", "programme", "program"}
    out, seen = [], set()
    for cells in rows:
        if all(c.lower() in headers for c in cells):
            continue
        date_obj, date_cell = None, None
        for c in cells:
            if not any(ch.isdigit() for ch in c):
                continue
            try:
                d1 = dateparser.parse(c, fuzzy=True, dayfirst=True, default=datetime(year, 1, 1))
                d2 = dateparser.parse(c, fuzzy=True, dayfirst=True, default=datetime(year, 6, 15))
            except Exception:
                continue
            # Only accept a fully-specified day+month (reject bare years like "2026")
            if d1.date() == d2.date():
                date_obj, date_cell = d1, c
                break
        if not date_obj:
            continue
        others = [c for c in cells if c != date_cell]
        title = max(others, key=len) if others else (cells[0].replace(date_cell, "").strip(" -–—:") or "Squadron event")
        d = date_obj.date()
        key = (d.isoformat(), title.lower())
        if key in seen:
            continue
        seen.add(key)
        out.append({"title": title[:200], "start": f"{d.isoformat()}T19:00:00",
                    "end": f"{d.isoformat()}T21:30:00", "description": "", "location": "Squadron HQ",
                    "date_label": d.strftime("%a %d %b %Y")})
    return out


@api_router.post("/events/import-docx")
async def import_docx(file: UploadFile = File(...), staff: dict = Depends(require_staff)):
    if not (file.filename or "").lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Please upload a Word .docx document.")
    data = await file.read()
    if len(data) > 10 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="File too large (maximum 10MB).")
    try:
        events = _parse_docx_events(data)
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Could not read the document: {exc}")
    return {"events": events, "count": len(events)}


@api_router.post("/events/import")
async def import_events(payload: EventsImport, staff: dict = Depends(require_staff)):
    created = 0
    for ie in payload.events:
        event = {"id": str(uuid.uuid4()), "title": ie.title, "description": ie.description,
                 "location": ie.location or "Squadron HQ", "start": ie.start, "end": ie.end,
                 "capacity": 0, "event_type": "standard", "participation": "attend", "points_value": 10,
                 "bids": [], "attendees": [], "created_by": staff["id"], "created_at": now_iso()}
        await db.events.insert_one(event)
        created += 1
    return {"created": created}


# ---------------------------------------------------------------------------
# Push subscriptions
# ---------------------------------------------------------------------------
class PushSubscription(BaseModel):
    subscription: dict
    model_config = ConfigDict(extra="ignore")


@api_router.get("/push/vapid-public-key")
async def vapid_public_key():
    return {"key": VAPID_PUBLIC_KEY, "enabled": bool(VAPID_PEM_PATH and VAPID_PUBLIC_KEY)}


@api_router.post("/push/subscribe")
async def push_subscribe(payload: PushSubscription, user: dict = Depends(get_current_user)):
    endpoint = payload.subscription.get("endpoint")
    if not endpoint:
        raise HTTPException(status_code=400, detail="Invalid subscription")
    await db.push_subscriptions.update_one(
        {"endpoint": endpoint},
        {"$set": {"endpoint": endpoint, "user_id": user["id"], "subscription": payload.subscription,
                  "updated_at": now_iso()}},
        upsert=True)
    return {"subscribed": True}


@api_router.post("/push/unsubscribe")
async def push_unsubscribe(payload: PushSubscription, user: dict = Depends(get_current_user)):
    endpoint = payload.subscription.get("endpoint")
    if endpoint:
        await db.push_subscriptions.delete_one({"endpoint": endpoint, "user_id": user["id"]})
    return {"unsubscribed": True}


@api_router.post("/push/test")
async def push_test(user: dict = Depends(get_current_user)):
    await push_to_user(user["id"], "1471 Horwich Squadron",
                       "Notifications are working \u2014 you'll be alerted here.", "/portal")
    return {"sent": True}


# ---------------------------------------------------------------------------
# DofE Diary
# ---------------------------------------------------------------------------

class DofEDiaryEntryCreate(BaseModel):
    section: str  # volunteering | skills | physical
    week_date: str  # YYYY-MM-DD (Monday of the week)
    content: str = Field(default="", max_length=10000)
    model_config = ConfigDict(extra="ignore")


class DofEDiaryEntryUpdate(BaseModel):
    content: Optional[str] = Field(default=None, max_length=10000)


DOFE_SECTIONS = {"volunteering", "skills", "physical"}


def _week_monday(d: date) -> str:
    return (d - timedelta(days=d.weekday())).isoformat()


@api_router.get("/dofe/diary/prompt-check")
async def dofe_diary_prompt_check(user: dict = Depends(get_current_user)):
    if user["role"] != "cadet":
        raise HTTPException(status_code=403, detail="Cadets only")
    today = date.today()
    weeks = [_week_monday(today - timedelta(weeks=i)) for i in range(4)]
    missing = []
    for week in weeks:
        for section in ["volunteering", "skills", "physical"]:
            entry = await db.dofe_diary.find_one(
                {"cadet_id": user["id"], "section": section, "week_date": week},
                {"_id": 0, "id": 1, "content": 1})
            if not entry or not entry.get("content", "").strip():
                missing.append({"week_date": week, "section": section})
    return {"missing": missing}


@api_router.get("/dofe/diary/export")
async def export_dofe_diary(cadet_id: Optional[str] = None, user: dict = Depends(get_current_user)):
    if user["role"] == "cadet":
        cid = user["id"]
    elif user["role"] in ("admin", "cfav"):
        if not cadet_id:
            raise HTTPException(status_code=400, detail="cadet_id required for staff access")
        cid = cadet_id
    elif user["role"] == "parent":
        linked = user.get("child_ids", []) or []
        if cadet_id:
            if cadet_id not in linked:
                raise HTTPException(status_code=403, detail="Forbidden")
            cid = cadet_id
        else:
            if len(linked) == 1:
                cid = linked[0]
            else:
                raise HTTPException(status_code=400, detail="cadet_id required for parent access")
    else:
        raise HTTPException(status_code=403, detail="Forbidden")
    cadet = await db.users.find_one({"id": cid}, {"_id": 0})
    if not cadet:
        raise HTTPException(status_code=404, detail="Cadet not found")
    entries = await db.dofe_diary.find({"cadet_id": cid}, {"_id": 0}).sort("week_date", 1).to_list(500)
    for e in entries:
        e["files"] = await db.dofe_diary_files.find(
            {"id": {"$in": e.get("file_ids", [])}}, {"_id": 0}).to_list(50)
    try:
        from reportlab.lib.pagesizes import A4 as _A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        import io as _io
    except Exception:
        raise HTTPException(status_code=500, detail="PDF export dependency missing")
    buf = _io.BytesIO()
    pdf_doc = SimpleDocTemplate(buf, pagesize=_A4,
                                rightMargin=2 * cm, leftMargin=2 * cm,
                                topMargin=2.5 * cm, bottomMargin=2 * cm)
    styles = getSampleStyleSheet()
    navy = colors.HexColor("#002F5F")
    red = colors.HexColor("#C80D30")
    title_s = ParagraphStyle("dt", parent=styles["Heading1"], textColor=navy, fontSize=16, spaceAfter=4)
    sub_s = ParagraphStyle("ds", parent=styles["Normal"], textColor=red, fontSize=11, spaceAfter=12)
    body_s = ParagraphStyle("db", parent=styles["Normal"], fontSize=9, spaceAfter=4, leading=14)
    week_s = ParagraphStyle("dw", parent=styles["Heading3"], textColor=navy, fontSize=10, spaceBefore=10, spaceAfter=4)
    sec_s = ParagraphStyle("dse", parent=styles["Heading4"], textColor=red, fontSize=9, spaceBefore=6, spaceAfter=2)
    empty_s = ParagraphStyle("de", parent=body_s, textColor=colors.grey)
    file_s = ParagraphStyle("df", parent=body_s, textColor=colors.grey, fontSize=8)
    cadet_name = f"{cadet.get('first_name', '')} {cadet.get('last_name', '')}".strip()
    story = [
        Paragraph("1471 Horwich Squadron RAF Air Cadets", title_s),
        Paragraph(f"DofE Diary \u2014 {(cadet.get('dofe_level') or 'Unknown').title()} Award", sub_s),
        Paragraph(f"Cadet: {cadet_name}", body_s),
        Paragraph(f"Generated: {datetime.now().strftime('%d %b %Y %H:%M')} UTC", body_s),
        HRFlowable(width="100%", thickness=2, color=navy, spaceAfter=12),
    ]
    from collections import defaultdict as _dd
    by_week = _dd(dict)
    for e in entries:
        by_week[e["week_date"]][e["section"]] = e
    for week_date in sorted(by_week.keys()):
        try:
            wd = date.fromisoformat(week_date)
            week_label = f"Week of {wd.strftime('%d %b %Y')}"
        except Exception:
            week_label = week_date
        story.append(Paragraph(week_label, week_s))
        for sec in ["volunteering", "skills", "physical"]:
            story.append(Paragraph(sec.title(), sec_s))
            e = by_week[week_date].get(sec)
            if e and e.get("content", "").strip():
                for para in e["content"].split("\n"):
                    if para.strip():
                        story.append(Paragraph(para.strip(), body_s))
            else:
                story.append(Paragraph("(No entry)", empty_s))
            if e and e.get("files"):
                story.append(Paragraph(
                    "Attachments: " + ", ".join(f["filename"] for f in e["files"]), file_s))
        story.append(Spacer(1, 6))
    pdf_doc.build(story)
    buf.seek(0)
    safe = cadet_name.replace(" ", "_") or "cadet"
    return Response(content=buf.read(), media_type="application/pdf",
                    headers={"Content-Disposition": f'attachment; filename="DofE_Diary_{safe}.pdf"'})


@api_router.get("/dofe/diary")
async def get_dofe_diary(cadet_id: Optional[str] = None, user: dict = Depends(get_current_user)):
    if user["role"] == "cadet":
        cid = user["id"]
    elif user["role"] in ("admin", "cfav"):
        if not cadet_id:
            raise HTTPException(status_code=400, detail="cadet_id required for staff access")
        cid = cadet_id
    elif user["role"] == "parent":
        linked = user.get("child_ids", []) or []
        if cadet_id:
            if cadet_id not in linked:
                raise HTTPException(status_code=403, detail="Forbidden")
            cid = cadet_id
        else:
            if len(linked) == 1:
                cid = linked[0]
            else:
                raise HTTPException(status_code=400, detail="cadet_id required for parent access")
    else:
        raise HTTPException(status_code=403, detail="Forbidden")
    entries = await db.dofe_diary.find({"cadet_id": cid}, {"_id": 0}).sort("week_date", -1).to_list(500)
    for e in entries:
        e["files"] = await db.dofe_diary_files.find(
            {"id": {"$in": e.get("file_ids", [])}}, {"_id": 0}).to_list(50)
    return entries


@api_router.post("/dofe/diary")
async def upsert_dofe_diary_entry(payload: DofEDiaryEntryCreate, user: dict = Depends(get_current_user)):
    if user["role"] != "cadet":
        raise HTTPException(status_code=403, detail="Cadets only")
    if payload.section not in DOFE_SECTIONS:
        raise HTTPException(status_code=400, detail="Invalid section")
    try:
        d = date.fromisoformat(payload.week_date)
        if d.weekday() != 0:
            raise ValueError
    except ValueError:
        raise HTTPException(status_code=400, detail="week_date must be a Monday in YYYY-MM-DD format")
    existing = await db.dofe_diary.find_one(
        {"cadet_id": user["id"], "section": payload.section, "week_date": payload.week_date},
        {"_id": 0})
    if existing:
        await db.dofe_diary.update_one(
            {"id": existing["id"]},
            {"$set": {"content": payload.content, "updated_at": now_iso()}})
        return {**existing, "content": payload.content, "updated_at": now_iso(),
                "files": await db.dofe_diary_files.find(
                    {"id": {"$in": existing.get("file_ids", [])}}, {"_id": 0}).to_list(50)}
    doc = {
        "id": str(uuid.uuid4()), "cadet_id": user["id"],
        "section": payload.section, "week_date": payload.week_date,
        "content": payload.content, "file_ids": [],
        "created_at": now_iso(), "updated_at": now_iso(),
    }
    await db.dofe_diary.insert_one(doc)
    return {**{k: v for k, v in doc.items() if k != "_id"}, "files": []}


@api_router.post("/dofe/diary/{entry_id}/upload")
async def upload_dofe_diary_file(
        entry_id: str, file: UploadFile = File(...), user: dict = Depends(get_current_user)):
    entry = await db.dofe_diary.find_one({"id": entry_id}, {"_id": 0})
    if not entry:
        raise HTTPException(status_code=404, detail="Entry not found")
    if entry["cadet_id"] != user["id"] and user["role"] not in ("admin", "cfav"):
        raise HTTPException(status_code=403, detail="Forbidden")
    data = await file.read()
    if len(data) > 15 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="File too large (maximum 15 MB).")
    fid = str(uuid.uuid4())
    gid = await fs.upload_from_stream(
        file.filename or "file", data,
        metadata={"dofe_file_id": fid, "content_type": file.content_type})
    fdoc = {
        "id": fid, "entry_id": entry_id,
        "filename": file.filename or "file",
        "content_type": file.content_type or "application/octet-stream",
        "size": len(data), "gridfs_id": str(gid),
        "uploaded_by": user["id"], "created_at": now_iso(),
    }
    await db.dofe_diary_files.insert_one(fdoc)
    await db.dofe_diary.update_one({"id": entry_id}, {"$push": {"file_ids": fid}})
    return {k: v for k, v in fdoc.items() if k != "_id"}


@api_router.get("/dofe/diary/files/{file_id}/download")
async def download_dofe_diary_file(file_id: str, user: dict = Depends(get_current_user)):
    fdoc = await db.dofe_diary_files.find_one({"id": file_id}, {"_id": 0})
    if not fdoc:
        raise HTTPException(status_code=404, detail="File not found")
    entry = await db.dofe_diary.find_one({"id": fdoc["entry_id"]}, {"_id": 0, "cadet_id": 1})
    if entry:
        if user["role"] == "cadet" and entry["cadet_id"] != user["id"]:
            raise HTTPException(status_code=403, detail="Forbidden")
        elif user["role"] in ("admin", "cfav"):
            pass
        elif user["role"] == "parent":
            linked = user.get("child_ids", []) or []
            if entry["cadet_id"] not in linked:
                raise HTTPException(status_code=403, detail="Forbidden")
        else:
            raise HTTPException(status_code=403, detail="Forbidden")
    stream = await fs.open_download_stream(ObjectId(fdoc["gridfs_id"]))
    data = await stream.read()
    return Response(content=data, media_type=fdoc["content_type"],
                    headers={"Content-Disposition": f'inline; filename="{fdoc["filename"]}"'})


@api_router.delete("/dofe/diary/{entry_id}/files/{file_id}")
async def delete_dofe_diary_file(
        entry_id: str, file_id: str, user: dict = Depends(get_current_user)):
    entry = await db.dofe_diary.find_one({"id": entry_id}, {"_id": 0})
    if not entry:
        raise HTTPException(status_code=404, detail="Entry not found")
    if entry["cadet_id"] != user["id"] and user["role"] not in ("admin",):
        raise HTTPException(status_code=403, detail="Forbidden")
    fdoc = await db.dofe_diary_files.find_one({"id": file_id, "entry_id": entry_id})
    if not fdoc:
        raise HTTPException(status_code=404, detail="File not found")
    try:
        await fs.delete(ObjectId(fdoc["gridfs_id"]))
    except Exception:
        pass
    await db.dofe_diary_files.delete_one({"id": file_id})
    await db.dofe_diary.update_one({"id": entry_id}, {"$pull": {"file_ids": file_id}})
    return {"deleted": True}


# ---------------------------------------------------------------------------
# Facebook Posts Integration
# ---------------------------------------------------------------------------

FACEBOOK_GRAPH_URL = "https://graph.facebook.com/v19.0"
FACEBOOK_PUBLIC_SCAN_URL = "https://r.jina.ai/http://www.facebook.com/1471HorwichRAFAC/"
_facebook_public_cache = {"at": None, "posts": []}


class FBConfig(BaseModel):
    page_id: str = Field(..., min_length=1)
    access_token: str = Field(..., min_length=1)


async def _get_fb_config() -> Optional[dict]:
    return await db.settings.find_one({"key": "facebook"}, {"_id": 0})


async def _do_fb_sync(config: dict) -> int:
    """Fetch public posts from the Graph API and upsert into fb_posts collection."""
    v = config.get("value", {})
    page_id = v.get("page_id", "")
    token = v.get("access_token", "")
    if not page_id or not token:
        return 0
    fields = "id,message,story,created_time,full_picture,permalink_url"
    async with httpx.AsyncClient(timeout=20.0) as client:
        resp = await client.get(
            f"{FACEBOOK_GRAPH_URL}/{page_id}/posts",
            params={"fields": fields, "access_token": token, "limit": 30},
        )
        resp.raise_for_status()
        data = resp.json()
    synced = 0
    for p in data.get("data", []):
        text = (p.get("message") or p.get("story") or "").strip()
        if not text:
            continue
        doc = {
            "fb_id": p["id"],
            "message": text,
            "created_time": p.get("created_time", ""),
            "full_picture": p.get("full_picture"),
            "permalink_url": p.get("permalink_url", ""),
            "synced_at": now_iso(),
        }
        await db.fb_posts.update_one({"fb_id": p["id"]}, {"$set": doc}, upsert=True)
        synced += 1
    await db.settings.update_one(
        {"key": "facebook"},
        {"$set": {"last_sync": now_iso()}},
    )
    return synced


def _clean_fb_link(url: str) -> str:
    link = (url or "").strip()
    if not link:
        return ""
    link = link.replace("http://www.facebook.com", "https://www.facebook.com")
    link = link.replace("http://facebook.com", "https://www.facebook.com")
    return link


def _headline_from_text(text: str) -> str:
    base = re.sub(r"\s+", " ", (text or "").strip())
    if not base:
        return "Facebook post"
    if len(base) <= 90:
        return base
    return f"{base[:87].rstrip()}..."


def _extract_public_fb_posts(markdown: str, limit: int) -> List[dict]:
    # The mirrored page is markdown-like content; each recent post appears as
    # "## **[Page Name](...)** ..." followed by a relative time link and media.
    blocks = markdown.split("\n## **[")
    out: List[dict] = []
    seen = set()
    for idx, raw in enumerate(blocks):
        if idx == 0:
            continue
        block = f"## **[{raw}"
        action_match = re.search(r"\*\*\s*(.+?)\n", block)
        action_text = (action_match.group(1) if action_match else "").strip()

        link_match = re.search(r"\n\[([^\]]+)\]\((https?://[^\)]+/posts/[^\)]+)\)", block)
        if not link_match:
            link_match = re.search(r"\n\[([^\]]+)\]\((https?://[^\)]+/photo/\?fbid=[^\)]+)\)", block)
        if not link_match:
            continue

        time_label = (link_match.group(1) or "").strip()
        permalink = _clean_fb_link(link_match.group(2))
        if not permalink or permalink in seen:
            continue
        seen.add(permalink)

        image_match = re.search(r"!\[[^\]]+\]\((https?://[^\)]+)\)", block)
        image_url = image_match.group(1).strip() if image_match else ""

        out.append({
            "fb_id": permalink.split("/")[-1][:80],
            "headline": _headline_from_text(action_text),
            "permalink_url": permalink,
            "full_picture": image_url,
            "created_time": "",
            "time_label": time_label,
        })
        if len(out) >= limit:
            break
    return out


async def _scan_public_fb_posts(limit: int = 5) -> List[dict]:
    now = datetime.now(timezone.utc)
    cached_at = _facebook_public_cache.get("at")
    cached_posts = _facebook_public_cache.get("posts") or []
    if cached_at and (now - cached_at).total_seconds() < 900 and cached_posts:
        return cached_posts[:limit]

    async with httpx.AsyncClient(timeout=25.0) as client:
        resp = await client.get(FACEBOOK_PUBLIC_SCAN_URL)
        resp.raise_for_status()
        md = resp.text

    posts = _extract_public_fb_posts(md, limit)
    if not posts:
        posts = [{
            "fb_id": "fallback-page-link",
            "headline": "See our latest posts on Facebook",
            "permalink_url": "https://www.facebook.com/1471HorwichRAFAC/",
            "full_picture": "",
            "created_time": "",
            "time_label": "",
        }]
    _facebook_public_cache["at"] = now
    _facebook_public_cache["posts"] = posts
    return posts[:limit]


@api_router.get("/facebook/config")
async def get_fb_config(admin: dict = Depends(require_roles("admin"))):
    doc = await _get_fb_config()
    count = await db.fb_posts.count_documents({})
    if not doc:
        return {"page_id": "", "has_token": False, "last_sync": None, "post_count": count}
    v = doc.get("value", {})
    return {
        "page_id": v.get("page_id", ""),
        "has_token": bool(v.get("access_token")),
        "last_sync": doc.get("last_sync"),
        "post_count": count,
    }


@api_router.put("/facebook/config")
async def update_fb_config(payload: FBConfig, admin: dict = Depends(require_roles("admin"))):
    await db.settings.update_one(
        {"key": "facebook"},
        {"$set": {
            "key": "facebook",
            "value": {"page_id": payload.page_id, "access_token": payload.access_token},
        }},
        upsert=True,
    )
    return {"saved": True}


@api_router.post("/facebook/sync")
async def sync_fb_posts(admin: dict = Depends(require_roles("admin"))):
    config = await _get_fb_config()
    if not config or not config.get("value", {}).get("access_token"):
        raise HTTPException(
            status_code=400,
            detail="Facebook not configured. Set your page ID and access token first.",
        )
    try:
        count = await _do_fb_sync(config)
        return {"synced": count, "ok": True}
    except httpx.HTTPStatusError as e:
        detail = f"Facebook API error {e.response.status_code}"
        try:
            body = e.response.json()
            msg = body.get("error", {}).get("message", "")
            if msg:
                detail += f": {msg}"
        except Exception:
            pass
        raise HTTPException(status_code=502, detail=detail)
    except httpx.RequestError as e:
        raise HTTPException(status_code=502, detail=f"Could not reach Facebook: {e}")


@api_router.get("/facebook/posts")
async def get_fb_posts(background_tasks: BackgroundTasks, limit: int = 20):
    """Public: return cached posts; auto-refresh in background if stale (>2 h)."""
    config = await _get_fb_config()
    if config and config.get("value", {}).get("access_token"):
        last = config.get("last_sync")
        needs = True
        if last:
            try:
                last_dt = datetime.fromisoformat(last.replace("Z", "+00:00"))
                needs = (datetime.now(timezone.utc) - last_dt).total_seconds() > 900
            except Exception:
                needs = True
        if needs:
            background_tasks.add_task(_do_fb_sync, config)
    posts = await db.fb_posts.find(
        {"message": {"$ne": ""}},
        {"_id": 0},
    ).sort("created_time", -1).to_list(min(limit, 5))
    return posts


@api_router.get("/facebook/public-scan")
async def get_fb_public_scan(limit: int = 5):
    lim = max(1, min(limit, 8))
    try:
        posts = await _scan_public_fb_posts(lim)
        return posts
    except Exception as e:
        logger.warning("Public Facebook scan failed: %s", e)
        return [{
            "fb_id": "fallback-page-link",
            "headline": "See our latest posts on Facebook",
            "permalink_url": "https://www.facebook.com/1471HorwichRAFAC/",
            "full_picture": "",
            "created_time": "",
            "time_label": "",
        }]


# ---------------------------------------------------------------------------
# Custom Activities (admin-managed, stored in MongoDB)
# ---------------------------------------------------------------------------

CUSTOM_ACTIVITY_ICONS = [
    "Plane", "Wind", "Mountain", "Award", "HeartPulse", "Compass", "Tent",
    "Trophy", "TentTree", "Shield", "BookOpen", "HeartHandshake", "GraduationCap",
    "Users", "Target", "Rocket", "Globe2", "Sparkles", "Star", "Briefcase",
    "Camera", "Music", "Wrench", "Flag", "Zap",
]


class CustomActivityCreate(BaseModel):
    slug: str = Field(..., min_length=1, pattern=r"^[a-z0-9-]+$")
    title: str = Field(..., min_length=1)
    strapline: str = Field(default="")
    text: str = Field(default="")
    long: List[str] = []
    highlights: List[str] = []
    quick_facts: List[str] = []
    what_to_expect: List[str] = []
    image_url: str = Field(default="")
    icon_name: str = Field(default="Compass")
    published: bool = True
    model_config = ConfigDict(extra="ignore")


class CustomActivityUpdate(BaseModel):
    title: Optional[str] = None
    strapline: Optional[str] = None
    text: Optional[str] = None
    long: Optional[List[str]] = None
    highlights: Optional[List[str]] = None
    quick_facts: Optional[List[str]] = None
    what_to_expect: Optional[List[str]] = None
    image_url: Optional[str] = None
    icon_name: Optional[str] = None
    published: Optional[bool] = None


@api_router.get("/activities/custom")
async def list_custom_activities(include_unpublished: bool = False,
                                  user: Optional[dict] = Depends(get_current_user) if False else None):
    q = {} if include_unpublished else {"published": True}
    rows = await db.custom_activities.find(q, {"_id": 0}).sort("title", 1).to_list(200)
    return rows


@api_router.get("/activities/custom/{slug}")
async def get_custom_activity(slug: str):
    act = await db.custom_activities.find_one({"slug": slug}, {"_id": 0})
    if not act:
        raise HTTPException(status_code=404, detail="Custom activity not found")
    return act


@api_router.post("/activities/custom")
async def create_custom_activity(
        payload: CustomActivityCreate,
        staff: dict = Depends(require_privileged_staff)):
    existing = await db.custom_activities.find_one({"slug": payload.slug})
    if existing:
        raise HTTPException(status_code=400, detail="An activity with this slug already exists")
    doc = {
        "id": str(uuid.uuid4()),
        **payload.model_dump(),
        "created_at": now_iso(),
        "updated_at": now_iso(),
        "created_by": staff["id"],
    }
    await db.custom_activities.insert_one(doc)
    return {k: v for k, v in doc.items() if k != "_id"}


@api_router.patch("/activities/custom/{slug}")
async def update_custom_activity(
        slug: str,
        payload: CustomActivityUpdate,
        staff: dict = Depends(require_privileged_staff)):
    updates = {k: v for k, v in payload.model_dump().items() if v is not None}
    if not updates:
        raise HTTPException(status_code=400, detail="No fields to update")
    updates["updated_at"] = now_iso()
    result = await db.custom_activities.find_one_and_update(
        {"slug": slug}, {"$set": updates},
        projection={"_id": 0}, return_document=True)
    if not result:
        raise HTTPException(status_code=404, detail="Custom activity not found")
    return result


@api_router.delete("/activities/custom/{slug}")
async def delete_custom_activity(
        slug: str,
        staff: dict = Depends(require_privileged_staff)):
    res = await db.custom_activities.delete_one({"slug": slug})
    if res.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Custom activity not found")
    return {"deleted": True}


app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

_progression_task = None


async def _progression_alert_loop():
    while True:
        try:
            await _send_classification_stagnation_alerts()
        except Exception as exc:  # pragma: no cover
            logger.error("Progression alert loop failed: %s", exc)
        await asyncio.sleep(12 * 60 * 60)


# ---------------------------------------------------------------------------
# Startup / seed
# ---------------------------------------------------------------------------
async def seed_users():
    demo = [
        {"email": ADMIN_EMAIL.lower(), "password": ADMIN_PASSWORD, "role": "admin",
         "first_name": "Squadron", "last_name": "Admin"},
        {"email": "cfav@1471horwich.org.uk", "password": "Cfav1471!", "role": "cfav",
         "first_name": "Alex", "last_name": "Instructor"},
        {"email": "cadet@1471horwich.org.uk", "password": "Cadet1471!", "role": "cadet",
         "first_name": "Sam", "last_name": "Cadet"},
        {"email": "parent@1471horwich.org.uk", "password": "Parent1471!", "role": "parent",
         "first_name": "Jordan", "last_name": "Parent"},
    ]
    ids = {}
    for d in demo:
        existing = await db.users.find_one({"email": d["email"]})
        if existing is None:
            uid = str(uuid.uuid4())
            await db.users.insert_one({
                "id": uid, "email": d["email"], "password_hash": hash_password(d["password"]),
                "role": d["role"], "first_name": d["first_name"], "last_name": d["last_name"],
                "login_username": await ensure_login_username(d["first_name"], d["last_name"]),
                "must_change_password": d["role"] == "cadet",
                "is_uniformed": True if d["role"] == "cfav" else None,
                "child_ids": [], "bonus_points": 0, "created_at": now_iso()})
            ids[d["role"]] = uid
        else:
            if d["role"] == "cfav" and "is_uniformed" not in existing:
                await db.users.update_one({"email": d["email"]}, {"$set": {"is_uniformed": True}})
            if not verify_password(d["password"], existing["password_hash"]):
                await db.users.update_one({"email": d["email"]},
                                          {"$set": {"password_hash": hash_password(d["password"])}})
            if not existing.get("login_username"):
                await db.users.update_one(
                    {"email": d["email"]},
                    {"$set": {"login_username": await ensure_login_username(d["first_name"], d["last_name"], existing.get("id"))}},
                )
            ids[d["role"]] = existing["id"]
    # Link demo parent to demo cadet
    if ids.get("parent") and ids.get("cadet"):
        await db.users.update_one({"id": ids["parent"]}, {"$set": {"child_ids": [ids["cadet"]]}})

    # Seed a welcome notice requiring acknowledgement (idempotent)
    if await db.notices.find_one({"title": "Welcome to the 1471 Members Area"}) is None:
        await db.notices.insert_one({
            "id": str(uuid.uuid4()),
            "title": "Welcome to the 1471 Members Area",
            "body": "Welcome! Please keep your contact details up to date and check the calendar regularly to bid for upcoming events.",
            "roles": ["cadet", "parent", "cfav"], "requires_ack": True,
            "created_by": "Squadron Admin", "created_at": now_iso()})


async def ensure_usernames_backfilled():
    missing = await db.users.find(
        {"$or": [{"login_username": {"$exists": False}}, {"login_username": ""}]},
        {"_id": 0, "id": 1, "first_name": 1, "last_name": 1},
    ).to_list(5000)
    for u in missing:
        username = await ensure_login_username(u.get("first_name", ""), u.get("last_name", ""), u.get("id"))
        await db.users.update_one({"id": u["id"]}, {"$set": {"login_username": username}})


@app.on_event("startup")
async def on_startup():
    global _progression_task
    await db.users.create_index("email", unique=True)
    await db.users.create_index("id", unique=True)
    await db.users.create_index("login_username", unique=True, sparse=True)
    await db.users.create_index([("role", 1), ("is_uniformed", 1)])
    await db.events.create_index("start")
    await db.notice_acks.create_index([("notice_id", 1), ("user_id", 1)], unique=True)
    await db.messages.create_index("member_id")
    await db.notifications.create_index([("user_id", 1), ("created_at", -1)])
    await db.documents.create_index([("visible_roles", 1), ("created_at", -1)])
    await db.learning_assignments.create_index([("cadet_ids", 1), ("created_at", -1)])
    await db.learning_submissions.create_index([("assignment_id", 1), ("cadet_id", 1)], unique=True)
    await db.learning_submissions.create_index([("cadet_id", 1), ("submitted_at", -1)])
    await db.push_subscriptions.create_index("endpoint", unique=True)
    await db.cadet_tracker.create_index("tracker_key", unique=True)
    await db.cadet_tracker.create_index("user_id")
    await db.cadet_progress_alerts.create_index("cadet_id", unique=True)
    await db.cfav_availability.create_index([("parade_date", 1), ("cfav_id", 1)])
    await db.cfav_event_ideas.create_index([("parade_date", 1), ("created_at", -1)])
    await db.cfav_skill_matrix.create_index("cfav_id", unique=True)
    await seed_users()
    await ensure_usernames_backfilled()
    _progression_task = asyncio.create_task(_progression_alert_loop())


@app.on_event("shutdown")
async def shutdown_db_client():
    global _progression_task
    if _progression_task:
        _progression_task.cancel()
    client.close()
